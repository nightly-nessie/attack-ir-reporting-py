#!/usr/bin/env python
# coding: utf-8
import json
import os
import re
import urllib.request
import uuid
import shutil
from datetime import datetime
import docx
from docx import Document
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.shared import RGBColor

template_directory = "templates"
resources_directory = "resources"
case_directory = str(uuid.uuid4())
parent_dir = os.path.join(os.getcwd())
template_path = os.path.join(parent_dir, template_directory)
resources_path = os.path.join(parent_dir, resources_directory)
case_path = os.path.join(parent_dir, case_directory)
os.mkdir(case_path, 0o744)
check_resources_path = os.path.isdir(resources_path)
dot_present = shutil.which("dot")
file_json_helper_enterprise_attack = os.path.join(resources_path, "helper_enterprise_attack.json")
file_json_helper_cis_controls_mapping = os.path.join(resources_path, "helper_cis_controls_mapping.json")
file_json_helper_nist_mapping = os.path.join(resources_path, "helper_nist_attack_mapping.json")
file_json_helper_ossem_mapping_array = os.path.join(resources_path, "helper_ossem_attack_mapping.json")
file_json_helper_atomicred_mapping_array = os.path.join(resources_path, "helper_atomicred_attack_mapping.json")
file_docx_template = os.path.join(template_path, "template.docx")
    
if not check_resources_path:
    os.mkdir(resources_path, 0o744)

def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    r = paragraph.add_run ()
    r._r.append (hyperlink)
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True
    return hyperlink

def process_text_with_links_code(text, paragraph):
    links_segments = re.split(r'(\[.*?\]\(.*?\))', text)
    for segment in links_segments:
        if re.match(r'\[.*?\]\(.*?\)', segment):
            definition, url = re.findall(r'\[(.*?)\]\((.*?)\)', segment)[0]
            hyperlink_text = f'{definition}'
            add_hyperlink(paragraph, hyperlink_text, url)
        else:
            code_segments = re.split(r'(<code>.*?</code>)', segment)
            for code_segment in code_segments:
                if re.match(r'<code>.*?</code>', code_segment):
                    code = re.findall(r'<code>(.*?)</code>', code_segment)[0]
                    paragraph.add_run(code).italic = True
                else:
                    paragraph.add_run(code_segment)
    return paragraph

def tactic_viz(tactic):
    filtered_array = [elem for elem in array_obj_sorted_construct if tactic in elem['attack_tactics']]
    x = 1
    tactic_column = ""
    tactic_column_attributes_list = []
    for item in filtered_array:
        attack_label = (item['attack_title']).replace(": ",":\\n")
        tactic_column_nodes = tactic.replace("-","")+str(x) + " [label = \"" + attack_label + "\"];"
        tactic_column = tactic_column + tactic_column_nodes
        tactic_column_attributes = tactic.replace("-","") + str(x)
        tactic_column_attributes_list.append(tactic_column_attributes)
        x = x + 1
    tactic_edges = " -> ".join(tactic_column_attributes_list)
    tactic_cluster = "subgraph cluster"+tactic.replace("-","")+" {label=\""+tactic.title().replace("-"," ")+"\";rank=same;style=dotted;" + tactic_edges + "};"
    tactic_content = (tactic_column + tactic_edges + ";" + tactic_cluster)
    return(tactic_content)

def get_attack_enterprise_json(attack_force, attack_version):
    """
    This function fetches the latest available ATT&CK(r) STIX JSON file from Github.
    A verification is performed whether the file already exists or not. No option is implemented to fetch a previous version of the file.

    :param force: bool, using this parameter will allow you to force a download of the ATT&CK(r) STIX JSON file from Github. Default value is False.
    :param version: str, using this parameter will allow you to select the version of the ATT&CK(r) STIX JSON file from Github. Default value is None.
    :return: dict, the contents of the JSON file
    """
    if attack_version:
        url_json_helper_enterprise_attack = f"https://raw.githubusercontent.com/mitre-attack/attack-stix-data/master/enterprise-attack/enterprise-attack-{attack_version}.json"
    else:
        url_json_helper_enterprise_attack = "https://raw.githubusercontent.com/mitre-attack/attack-stix-data/master/enterprise-attack/enterprise-attack.json"
    print("\u2139 The ATT&CK\u00AE JSON STIX file is required to continue. It will be downloaded if not already present in the folder")
    if not os.path.isfile(file_json_helper_enterprise_attack) or attack_force or attack_version:
        urllib.request.urlretrieve(url_json_helper_enterprise_attack, file_json_helper_enterprise_attack)
        print(f"{url_json_helper_enterprise_attack} has been downloaded.")
    else:
        with open(file_json_helper_enterprise_attack, 'r', encoding='utf-8') as f:
            array_obj_complete_attack = json.load(f)
            modified = array_obj_complete_attack["objects"][0]["modified"]
            modified_dt = datetime.strptime(modified, "%Y-%m-%dT%H:%M:%S.%fZ")
            modified_str = modified_dt.strftime("%F")
            version = array_obj_complete_attack["objects"][0]["x_mitre_version"]
        print(f"\u2139 The local ATT&CK\u00AE JSON STIX file was present already and was last modified on {modified_str}. It serves MITRE ATT&CK\u00AE version {version}")
        print("\u2139 Consider running 'get_attack_enterprise_json(force=True)' to fetch the latest version or run 'get_attack_enterprise_json(attack_version='11.0')' for a specific version. The current file is not overwritten.")

def get_cis_controls_json(cis_force):
    print("\u2139 The CIS Controls ATT&CK\u00AE mapping JSON STIX file is required to continue. It will be silently downloaded if not already present in the folder")
    if not os.path.exists(file_json_helper_cis_controls_mapping) or cis_force:
        url_json_helper_cis_controls_mapping = "https://raw.githubusercontent.com/nightly-nessie/attack-cis-controls/main/cis-controls-8-enterprise-attack-12.json"
        urllib.request.urlretrieve(url_json_helper_cis_controls_mapping, file_json_helper_cis_controls_mapping)

def get_nist_controls_json(nist_force):
    """
    This function downloads the NIST 800-53 Rev 5 Controls ATT&CK mapping JSON STIX file if it is not already present in the folder.
    """
    print("\u2139 The NIST 800-53 Rev 5 Controls ATT&CK® mapping JSON STIX file is required to continue. It will be silently downloaded if not already present in the folder.")
    if not os.path.isfile(file_json_helper_nist_mapping) or nist_force:
        url_json_helper_nist_mapping = "https://raw.githubusercontent.com/center-for-threat-informed-defense/attack-control-framework-mappings/main/frameworks/attack_10_1/nist800_53_r5/stix/nist800-53-r5-enterprise-attack.json"
        urllib.request.urlretrieve(url_json_helper_nist_mapping, file_json_helper_nist_mapping)

def get_ossem_json(ossem_force):
    """
    This function downloads the OSSEM ATT&CK mapping JSON file if it is not already present in the folder.
    """
    print("\u2139 The OSSEM ATT&CK mapping JSON file is required to continue. It will be silently downloaded if not already present in the folder.")
    if not os.path.exists(file_json_helper_ossem_mapping_array) or ossem_force:
        url_json_helper_ossem_mapping = "https://raw.githubusercontent.com/OTRF/OSSEM-DM/main/use-cases/mitre_attack/techniques_to_events_mapping.json"
        urllib.request.urlretrieve(url_json_helper_ossem_mapping, file_json_helper_ossem_mapping_array)

def get_atomic_red_team_json(atomicred_force):
    """
    This function downloads the Red Canary Atomic Red Team tests mapping JSON file if it is not already present in the folder.
    The downloaded file is used to get the complete atomic red team mapping array.
    """
    print("\u2139 The Red Canary Atomic Red Team tests mapping JSON file is required to continue. It will be silently downloaded if not already present in the folder")
    if not os.path.exists(file_json_helper_atomicred_mapping_array) or atomicred_force:
        url_json_helper_atomicred_mapping = "https://raw.githubusercontent.com/redcanaryco/atomic-red-team/master/atomics/Indexes/Attack-Navigator-Layers/art-navigator-layer.json"
        urllib.request.urlretrieve(url_json_helper_atomicred_mapping, file_json_helper_atomicred_mapping_array)

def get_resources(attack_force=False, attack_version=None, cis_force=False, nist_force=False, ossem_force=False, atomicred_force=False):
    get_attack_enterprise_json(attack_force, attack_version)
    get_cis_controls_json(cis_force)
    get_nist_controls_json(nist_force)
    get_ossem_json(ossem_force)
    get_atomic_red_team_json(atomicred_force)

def get_resources_content():
    with open(file_json_helper_enterprise_attack, 'r', encoding='utf-8') as f:
        array_obj_complete_attack = json.load(f)
    with open(file_json_helper_cis_controls_mapping, 'r', encoding='utf-8') as f:
        array_obj_complete_cis_controls_mapping = json.load(f)
    with open(file_json_helper_nist_mapping, 'r', encoding='utf-8') as f:
        array_obj_complete_nist_mapping = json.load(f)
    with open(file_json_helper_ossem_mapping_array, 'r', encoding='utf-8') as f:
        array_obj_complete_ossem_mapping = json.load(f)
    with open(file_json_helper_atomicred_mapping_array, 'r', encoding='utf-8') as f:
        array_obj_complete_atomicred_mapping = json.load(f)
    globals()["array_obj_complete_attack"] = array_obj_complete_attack
    globals()["array_obj_complete_cis_controls_mapping"] = array_obj_complete_cis_controls_mapping
    globals()["array_obj_complete_nist_mapping"] = array_obj_complete_nist_mapping
    globals()["array_obj_complete_ossem_mapping"] = array_obj_complete_ossem_mapping
    globals()["array_obj_complete_atomicred_mapping"] = array_obj_complete_atomicred_mapping
    
def set_attack_empty(list_obj_attack_techniques=None):
    get_resources_content()
    filtered_objects = [obj for obj in array_obj_complete_attack["objects"] if obj.get('x_mitre_deprecated') != True and obj.get('revoked') != True and obj.get('type') == 'attack-pattern']
    filtered_external_references = [ref for obj in filtered_objects for ref in obj.get('external_references', []) if ref.get('source_name') == 'mitre-attack']
    list_obj_complete_techniques = [{'attack_id': ref.get('external_id')} for ref in filtered_external_references]
    if list_obj_attack_techniques is None:
        while True:
            validated = True
            list_obj_attack_techniques = input("🔨 Give a single or a semicolon separated list of ATT&CK® IDs (for example: T1566.002;T1018;T1033): ").split(";")
            list_obj_selected_attack_techniques = [{'attack_id': x} for x in list_obj_attack_techniques]
            for instance in list_obj_selected_attack_techniques:
                # Check if the item exists in the other list
                if not any(current.get("attack_id") == instance["attack_id"] for current in list_obj_complete_techniques):
                    # If it doesn't exist, print a message and set the validated flag to False
                    print("⚠️ "+ instance["attack_id"] + " does not exist in the current ATT&CK® Enterprise JSON. Please verify your input.")
                    validated = False
                    break
            if validated:
                globals()["list_obj_selected_attack_techniques"] = list_obj_selected_attack_techniques
                break
    else:
        list_obj_attack_techniques = list_obj_attack_techniques.split(";")
        while True:
            validated = True
            list_obj_selected_attack_techniques = [{'attack_id': x} for x in list_obj_attack_techniques]
            for instance in list_obj_selected_attack_techniques:
                # Check if the item exists in the other list
                if not any(current.get("attack_id") == instance["attack_id"] for current in list_obj_complete_techniques):
                    # If it doesn't exist, print a message and set the validated flag to False
                    print("⚠️ "+ instance["attack_id"] + " does not exist in the current ATT&CK® Enterprise JSON. Please verify your input.")
                    list_obj_attack_techniques = input("🔨 Give a single or a semicolon separated list of ATT&CK® IDs (for example: T1566.002;T1018;T1033): ").split(";")
                    validated = False
                    break
            if validated:
                globals()["list_obj_selected_attack_techniques"] = list_obj_selected_attack_techniques
                break
    new_attackconstruct()

def new_attackconstruct():
    array_obj_complete_mapping_external_id_attack_pattern = list(filter(lambda x: x['type'] == 'attack-pattern', array_obj_complete_attack["objects"]))
    array_obj_complete_mapping_external_id_attack_pattern = list(map(lambda x: {'id': x['id'], **x['external_references'][0]}, array_obj_complete_mapping_external_id_attack_pattern))
    array_obj_complete_mapping_external_id_attack_pattern = list(filter(lambda x: x['source_name'] == 'mitre-attack', array_obj_complete_mapping_external_id_attack_pattern))
    array_obj_complete_mapping_external_id_attack_pattern = list(map(lambda x: {'external_id': x['external_id'], 'id': x['id']}, array_obj_complete_mapping_external_id_attack_pattern))
    array_obj_filtered_mapping_external_id_attack_pattern = list(filter(lambda x: x.get('external_id') in [d.get('attack_id') for d in list_obj_selected_attack_techniques], array_obj_complete_mapping_external_id_attack_pattern))
    array_obj_sorted_mapping_external_id_attack_pattern = sorted(array_obj_filtered_mapping_external_id_attack_pattern, key=lambda x: x['external_id'])
    array_obj_complete_construct = []
    for attack_id in array_obj_sorted_mapping_external_id_attack_pattern:
        obj_filtered_attack_attack_pattern = next((obj for obj in array_obj_complete_attack["objects"] if obj["type"] == "attack-pattern" and obj["id"] == attack_id["id"]), None)
        content_introduction_attack_name = obj_filtered_attack_attack_pattern["name"]
        array_obj_complete_attack_tactics = []
        for phase_name in obj_filtered_attack_attack_pattern["kill_chain_phases"]:
            array_obj_complete_attack_tactics.append(phase_name["phase_name"])
        content_introduction_attack_description = obj_filtered_attack_attack_pattern["description"]
        content_introduction_attack_description = re.sub(r'\(Citation:.*\)', '', content_introduction_attack_description)
        content_introduction_attack_description = re.sub(r"\r?\n\r?\n", "\n", content_introduction_attack_description)
        obj_filtered_attack_attack_pattern_property_external_references = next((ref for ref in obj_filtered_attack_attack_pattern["external_references"] if ref["source_name"] == "mitre-attack"), None)
        content_introduction_attack_url = obj_filtered_attack_attack_pattern_property_external_references["url"]
        obj_filtered_attack_attack_pattern_property_external_id = obj_filtered_attack_attack_pattern_property_external_references["external_id"]
        guid = str(uuid.uuid4())
        dict_row = {
            "attack_title": obj_filtered_attack_attack_pattern_property_external_id + ": " + content_introduction_attack_name,
            "attack_name": content_introduction_attack_name,
            "attack_id": obj_filtered_attack_attack_pattern_property_external_id,
            "attack_tactics": array_obj_complete_attack_tactics,
            "attack_all_tactics": array_obj_complete_attack_tactics,
            "attack_url": content_introduction_attack_url,
            "attack_description": content_introduction_attack_description,
            "guid": guid
        }
        array_obj_complete_construct.append(dict_row)
    array_obj_selected_construct = []
    for attack in array_obj_complete_construct:
        if (len(attack["attack_tactics"])) == 1:
            array_obj_selected_construct.append(attack)
        else:
            print("\nMultiple tactics were found for " + str(attack["attack_id"]) + ": " + (", ".join((attack["attack_all_tactics"])).replace("-", " ")).title())
            for tactic in attack["attack_tactics"]:
                guid = str(uuid.uuid4())
                split_tactic = attack.copy()
                split_tactic["attack_tactics"] = [tactic]
                split_tactic["guid"] = guid
                beautyfy_split_tactic = str(split_tactic["attack_tactics"][0])
                beautyfy_split_tactic = (beautyfy_split_tactic.replace("-", " ")).title()
                query_add_tactic = input("\u2328 Do you want to add " + str(split_tactic["attack_title"]) + "/" + beautyfy_split_tactic + " pair ([Y]/N) ")
                if query_add_tactic == 'y' or query_add_tactic == "Y" or not query_add_tactic:
                    array_obj_selected_construct.append(split_tactic)
                    print("\u2328 " + str(split_tactic["attack_title"]) + "/" + beautyfy_split_tactic + " pair is added.")
                else:
                    pass
    attack_tactic_ranks = {
        'initial-access': 0,
        'execution': 1,
        'persistence': 2,
        'privilege-escalation': 3,
        'defense-evasion': 4,
        'credential-access': 5,
        'discovery': 6,
        'lateral-movement': 7,
        'collection': 8,
        'command-and-control': 9,
        'exfiltration': 10,
        'impact': 11
    }
    array_obj_sorted_construct = sorted(array_obj_selected_construct, key=lambda x: attack_tactic_ranks[x['attack_tactics'][0]])
    globals()["array_obj_sorted_construct"] = array_obj_sorted_construct
    globals()["array_obj_filtered_mapping_external_id_attack_pattern"] = array_obj_filtered_mapping_external_id_attack_pattern
    globals()["attack_tactic_ranks"] = attack_tactic_ranks       

def new_condensed_navigator():
    unique_attack_tactics = set()
    for dictionary in array_obj_sorted_construct:
        unique_attack_tactics.update(dictionary['attack_tactics'])
    unique_attack_tactics = list(unique_attack_tactics)
    sorted_unique_attack_tactic = sorted(unique_attack_tactics, key=lambda x: attack_tactic_ranks.get(x))
    navigator_header_viz = "digraph customer {layout=dot;label = \"\";labelloc = \"t\";node [style=rounded shape=Mrecord style=filled fillcolor = lightgrey color = lightgrey];edge [style=\"invis\"];"
    navigator_content_viz = ""
    for tactic in sorted_unique_attack_tactic:
        navigator_content_viz += tactic_viz(tactic)
    condensed_navigator_graphviz = navigator_header_viz + navigator_content_viz + "}"
    file_condensed_navigator_dot = os.path.join(case_path, "condensed_navigator.dot")
    with open(file_condensed_navigator_dot, 'w') as file_graph_dot:
        file_graph_dot.write(condensed_navigator_graphviz)
    from subprocess import check_call
    file_condensed_navigator = os.path.join(case_path, "condensed_navigator.png")
    check_call(['dot','-Tpng',file_condensed_navigator_dot,'-o',file_condensed_navigator,'-Gsize=5,3\!','-Gdpi=300'])
    globals()["file_condensed_navigator"] = file_condensed_navigator

def new_attackdocintroduction():
    if dot_present is not None:
        new_condensed_navigator()
    else:
       pass
    file_docx_introduction = os.path.join(case_path, document_prefix + "introduction.docx")
    document = Document(file_docx_template)
    document.add_heading("Introduction",1)
    document.add_paragraph("This annex describes the possible mitigations, controls and eventually detections to implement to avoid a similar incident from happening again. The identified adversary TTPs (Techniques, Procedures and Tactics) are the result from the investigation conducted by CPIRT. The information presented stems from the common library for adversarial TTPs, the MITRE ATT&CK® Framework [https://attack.mitre.org/]. The different techniques are listed, explained, and linked with the adversary tactics. Tactics are the goals an adversary wants to achieve. Next, based on these techniques, possible mitigations are listed, each with a description and relation with both the MITRE ATT&CK® Techniques and CIS Controls. Some environments do not allow or struggle implementing the presented mitigations/controls. To cover these gaps, detections should be put in place. Coverage of the possible detections against the identified Techniques also includes the platform (IaaS, Containers, Linux, Windows ...) and the collection layer (Network, Host ...) to deploy the detection. Some detections may not be relevant for the environment as the platform may not be in use. The indication of the platform makes it straightforward to disregard those irrelevant detections.")
    document.add_heading("Techniques",1)
    document.add_paragraph("According to the MITRE ATT&CK® Framework 'Techniques' represent 'how' an adversary achieves a tactical goal (tactic) by performing an action. For example, an adversary may dump credentials to achieve credential access. Below are the identified MITRE ATT&CK® Techniques listed which provide insight in the actions performed by perpetrators during this incident. Depending on the available information and artefacts, this may not be an exhaustive list but should provide a very reasonable starting point to understand the techniques used and the follow up mitigations/controls to implement. Assure you have put detections in place where mitigations/controls were not implemented or are insufficient.")
    for item in array_obj_sorted_construct:
        document.add_heading(item['attack_title'],2)
        table = document.add_table(rows=0,cols=1)
        row_cells = table.add_row().cells
        attack_tactic = (item['attack_tactics'][0]).replace("-", " ")
        row_cells[0].text = "Selected ATT&CK® Tactic: " + (attack_tactic).title()
        table.add_row()
        row_cells = table.add_row().cells
        attackurl = row_cells[0].paragraphs[0]
        add_hyperlink(attackurl, "ATT&CK® URL: " + item['attack_id'], item['attack_url'])
        row_cells = table.add_row().cells
        para = row_cells[0].add_paragraph()
        text = item['attack_description']
        process_text_with_links_code(text, para)
    if dot_present is not None:
        document.add_page_break()
        document.add_picture(file_condensed_navigator)
    else:
       pass
    document.save(file_docx_introduction)

def new_attackmitigationsconstruct():
    array_obj_complete_attack_mitigations = [obj for obj in array_obj_complete_attack["objects"] if obj.get("relationship_type") == "mitigates"]
    array_obj_filtered_attack_mitigations = [obj for obj in array_obj_complete_attack_mitigations if obj["target_ref"] in [attack_pattern["id"] for attack_pattern in array_obj_filtered_mapping_external_id_attack_pattern] and obj.get("x_mitre_deprecated") != True]
    array_obj_complete_mitigations = []
    array_obj_filtered_cis_controls_prio = []
    for mitigation in array_obj_filtered_attack_mitigations:
        obj_course_of_action_property_guid = mitigation['source_ref']
        obj_mitigation = next((obj for obj in array_obj_complete_attack['objects'] if obj['type'] == 'course-of-action' and obj['id'] == obj_course_of_action_property_guid), None)
        if obj_mitigation and obj_mitigation.get('x_mitre_deprecated', False) == True:
            pass
        else:
            obj_course_of_action_property_guid = mitigation['source_ref']
            obj_mitigation_property_description = mitigation['description']
            obj_mitigation_property_description = re.sub(r'\(Citation:.*\)', '', obj_mitigation_property_description)
            obj_mitigation_property_description_clean = re.sub(r"\r?\n\r?\n", "`n", obj_mitigation_property_description)
            obj_mitigation_attack_pattern = next((obj for obj in array_obj_filtered_mapping_external_id_attack_pattern if obj['id'] == mitigation['target_ref']), None)
            mitigation_component_block = next((obj for obj in array_obj_complete_attack['objects'] if obj['type'] == 'course-of-action' and obj['id'] == obj_course_of_action_property_guid), None)
            obj_mitigation_property_id = next((ref for ref in mitigation_component_block['external_references'] if ref['source_name'] == 'mitre-attack'), None)
            array_obj_filtered_cis_controls_mapping = [obj for obj in array_obj_complete_cis_controls_mapping['objects'] if obj.get('target_ref', '') == mitigation['source_ref']]
            array_obj_complete_cis_control_content = []
            for mapping in array_obj_filtered_cis_controls_mapping:
                obj_complete_cis_control = [obj for obj in array_obj_complete_cis_controls_mapping['objects'] if obj['type'] == 'course-of-action']
                obj_complete_cis_control = [obj for obj in obj_complete_cis_control if obj.get('id', '') == mapping['source_ref']]
                content_cis_controls = obj_complete_cis_control[0]["external_references"][0]["external_id"] + " " + obj_complete_cis_control[0]["name"]
                array_obj_complete_cis_control_content.append(content_cis_controls)
            for mapping in array_obj_filtered_cis_controls_mapping:
                obj_complete_cis_control = [obj for obj in array_obj_complete_cis_controls_mapping['objects'] if obj['type'] == 'course-of-action']
                obj_complete_cis_control = [obj for obj in obj_complete_cis_control if obj.get('id', '') == mapping['source_ref']]
                array_cis_row = {
                    "cis_control_id": obj_complete_cis_control[0].get("external_references")[0].get("external_id"),
                    "cis_control_name": obj_complete_cis_control[0].get("name"),
                    "cis_control_ig": obj_complete_cis_control[0].get("x_cis_ig")
                }
                array_obj_filtered_cis_controls_prio.append(array_cis_row)
            query_content_cis_controls = not bool(array_obj_complete_cis_control_content)
            if not query_content_cis_controls:
                content_cis_controls_body = "\n".join(array_obj_complete_cis_control_content)
            else:
                content_cis_controls_body = "There is no CIS Control® mapped with this Mitigation."
            nist_coas = [obj for obj in array_obj_complete_nist_mapping["objects"] if obj.get("relationship_type") == "mitigates"]
            nist_coas = [obj for obj in nist_coas if obj.get('target_ref', '') == mitigation['target_ref']]
            nist_control_array = []
            for coa in nist_coas:
                nist_coa_block = [obj for obj in array_obj_complete_nist_mapping["objects"] if obj.get("type") == "course-of-action"]
                nist_coa_block = [obj for obj in nist_coa_block if obj.get('id', '') == coa['source_ref']]
                nist_control_id = nist_coa_block[0]["external_references"][0]["external_id"]
                nist_control_name = nist_coa_block[0]["name"]
                nist_string = nist_control_id + " " + nist_control_name
                nist_control_array.append(nist_string)
            nist_control_array.sort()
            nist_control_body = "\n".join(nist_control_array)
            array_mitigations_row = {
                "name": mitigation_component_block["name"],
                "external_id": obj_mitigation_property_id["external_id"],
                "url": obj_mitigation_property_id["url"],
                "description": obj_mitigation_property_description_clean,
                "attack_id": obj_mitigation_attack_pattern["external_id"],
                "cis_control": content_cis_controls_body,
                "nist_control": nist_control_body
            }
            array_obj_complete_mitigations.append(array_mitigations_row)
    array_obj_sorted_mitigations = sorted(array_obj_complete_mitigations, key=lambda x: x.get('external_id', ''))
    from collections import defaultdict
    grouped_cis_controls = defaultdict(list)
    for control in array_obj_filtered_cis_controls_prio:
        grouped_cis_controls[control['cis_control_id']].append(control)
    array_obj_complete_cis_controls_prio = []
    for control_id, controls in grouped_cis_controls.items():
        cis_control_name = controls[0]['cis_control_name']
        cis_control_ig = controls[0]['cis_control_ig']
        cis_control_count = len(controls)
        new_obj = {
            'cis_control_id': control_id,
            'cis_control_name': cis_control_name,
            'cis_control_ig': cis_control_ig,
            'cis_control_count': cis_control_count
        }
        array_obj_complete_cis_controls_prio.append(new_obj)
    array_obj_complete_cis_controls_prio_sorted = sorted(array_obj_complete_cis_controls_prio, key=lambda x: (x['cis_control_ig'], -x['cis_control_count'], x['cis_control_id']))
    globals()["array_obj_sorted_mitigations"] = array_obj_sorted_mitigations
    globals()["array_obj_complete_cis_controls_prio_sorted"] = array_obj_complete_cis_controls_prio_sorted

def get_attackmitigationsmappings(ciscontrols,nistcontrols):
    if (ciscontrols and not nistcontrols):
        switch_control_mapping_selection = "CX"
    if (nistcontrols and not ciscontrols):
        switch_control_mapping_selection = "XN"
    if (ciscontrols and nistcontrols):
        switch_control_mapping_selection = "CN"
    if (not nistcontrols and not ciscontrols):
        query_cis_controls_mapping = input("\u2328 Do you want to generate the CIS Controls® v8 mapping? ([Y]/N) ")
        if not query_cis_controls_mapping or query_cis_controls_mapping.upper() == "Y":
            query_nist_controls_mapping = input("\u2328 Do you want to generate the NIST 800-53 Rev 5 Controls mapping? (Y/[N]) ")
            if not query_nist_controls_mapping or query_nist_controls_mapping.upper() == "N":
                switch_control_mapping_selection = "CX"
            else:
                switch_control_mapping_selection = "CN"
        else:
            query_nist_controls_mapping = input("\u2328 Do you want to generate the NIST 800-53 Rev 5 Controls mapping? (Y/[N]) ")
            if not query_nist_controls_mapping or query_nist_controls_mapping.upper() == "N":
                switch_control_mapping_selection = "XX"
            else:
                switch_control_mapping_selection = "XN"
    globals()["switch_control_mapping_selection"] = switch_control_mapping_selection

def new_attackdocmitigations(ciscontrols,nistcontrols):
    new_attackmitigationsconstruct()
    get_attackmitigationsmappings(ciscontrols,nistcontrols)
    file_docx_mitigations = os.path.join(case_path, document_prefix + "mitigations.docx")
    if switch_control_mapping_selection == "CN":
        document = Document(file_docx_template)
        document.add_heading("Mitigations/Controls",1)
        document.add_paragraph("Mitigations represent security concepts and classes of technologies that can be used to prevent (Sub)-Techniques from being successfully executed.")
        document.add_paragraph()
        document.add_heading("Mitigations Resume",2)
        document.add_paragraph()
        for mitigation in array_obj_sorted_mitigations:
            bulleted = document.add_paragraph(style='List Bullet')
            bulleted.add_run(mitigation["description"])
        document.add_page_break()  
        document.add_heading("Mitigations Overview",2)
        document.add_paragraph()
        document.add_paragraph("The mitigations listed below are mapped with the CIS Controls® v8 and the NIST 800-53 Rev 5 Controls. This mapping demonstrates which Controls are supported with the implementation of the corresponding Mitigations.")
        table_mitigations = document.add_table(rows=0,cols=3)
        table_mitigations.style = 'Table Grid'
        row_cells = table_mitigations.add_row().cells
        row_cells[0].paragraphs[0].add_run('Mitigation ID: Name').bold = True
        row_cells[1].paragraphs[0].add_run('Mitigation URL').bold = True
        row_cells[2].paragraphs[0].add_run('Covered ATT&CK® Technique').bold = True
        row_cells = table_mitigations.add_row().cells
        row_cells[0].merge(row_cells[2])
        row_cells[0].paragraphs[0].add_run('Description').bold = True
        row_cells = table_mitigations.add_row().cells
        row_cells[0].merge(row_cells[2])
        row_cells[0].paragraphs[0].add_run('CIS Controls® v8').bold = True
        row_cells = table_mitigations.add_row().cells
        row_cells[0].merge(row_cells[2])
        row_cells[0].paragraphs[0].add_run('NIST 800-53 Rev 5 Controls').bold = True
        for mitigation in array_obj_sorted_mitigations:
            row_cells = table_mitigations.add_row().cells
            row_cells[0].paragraphs[0].add_run(mitigation["external_id"] + ": " + mitigation["name"]).bold = True
            mitigationurl = row_cells[1].paragraphs[0]
            add_hyperlink(mitigationurl,mitigation['external_id'],mitigation['url'])
            techniqueurl = row_cells[2].paragraphs[0]
            content_generated_url = mitigation["attack_id"].replace(".", "/")
            add_hyperlink(techniqueurl,mitigation['attack_id'],"https://attack.mitre.org/techniques/" + content_generated_url)
            row_cells = table_mitigations.add_row().cells
            row_cells[0].merge(row_cells[2])
            row_cells[0].text = mitigation["description"]
            row_cells = table_mitigations.add_row().cells
            row_cells[0].merge(row_cells[2])
            row_cells[0].text = mitigation["cis_control"]
            row_cells = table_mitigations.add_row().cells
            row_cells[0].merge(row_cells[2])
            row_cells[0].text = mitigation["nist_control"]
        document.add_page_break()  
        document.add_heading("CIS Controls® Implementation Priority Guideline",2)
        document.add_paragraph("Below list presents a possible implementation priority, based on the lowest implementation groups where the CIS Control® is associated with and the weight of that specific CIS Control® in the mapping with the identified ATT&CK® (Sub-)Techniques and their associated Mitigations.")
        table_cis_controls_prio = document.add_table(rows=0,cols=4)
        table_cis_controls_prio.style = 'Table Grid'
        row_cells = table_cis_controls_prio.add_row().cells
        row_cells[0].paragraphs[0].add_run('Control® ID').bold = True
        row_cells[1].paragraphs[0].add_run('Control® Description').bold = True
        row_cells[2].paragraphs[0].add_run('IG').bold = True
        row_cells[3].paragraphs[0].add_run('Relative Weight').bold = True
        for mitigation in array_obj_complete_cis_controls_prio_sorted:
            row_cells = table_cis_controls_prio.add_row().cells
            row_cells[0].paragraphs[0].add_run(mitigation["cis_control_id"]).bold = True
            row_cells[1].text = mitigation["cis_control_name"]
            row_cells[2].text = mitigation["cis_control_ig"]
            row_cells[3].text = (str(mitigation["cis_control_count"]))
        document.save(file_docx_mitigations)
    elif switch_control_mapping_selection == "XN":
        document = Document(file_docx_template)
        document.add_heading("Mitigations/Controls",1)
        document.add_paragraph("Mitigations represent security concepts and classes of technologies that can be used to prevent (Sub)-Techniques from being successfully executed.")
        document.add_paragraph()
        document.add_heading("Mitigations Resume",2)
        document.add_paragraph()
        for mitigation in array_obj_sorted_mitigations:
            bulleted = document.add_paragraph(style='List Bullet')
            bulleted.add_run(mitigation["description"])
        document.add_page_break()  
        document.add_heading("Mitigations Overview",2)
        document.add_paragraph()
        document.add_paragraph("The mitigations listed below are mapped with the NIST 800-53 Rev 5 Controls. This mapping demonstrates which Controls are supported with the implementation of the corresponding Mitigations.")
        table_mitigations = document.add_table(rows=0,cols=3)
        table_mitigations.style = 'Table Grid'
        row_cells = table_mitigations.add_row().cells
        row_cells[0].paragraphs[0].add_run('Mitigation ID: Name').bold = True
        row_cells[1].paragraphs[0].add_run('Mitigation URL').bold = True
        row_cells[2].paragraphs[0].add_run('Covered ATT&CK® Technique').bold = True
        row_cells = table_mitigations.add_row().cells
        row_cells[0].merge(row_cells[2])
        row_cells[0].paragraphs[0].add_run('Description').bold = True
        row_cells = table_mitigations.add_row().cells
        row_cells[0].merge(row_cells[2])
        row_cells[0].paragraphs[0].add_run('NIST 800-53 Rev 5 Controls').bold = True
        for mitigation in array_obj_sorted_mitigations:
            row_cells = table_mitigations.add_row().cells
            row_cells[0].paragraphs[0].add_run(mitigation["external_id"] + ": " + mitigation["name"]).bold = True
            mitigationurl = row_cells[1].paragraphs[0]
            add_hyperlink(mitigationurl,mitigation['external_id'],mitigation['url'])
            techniqueurl = row_cells[2].paragraphs[0]
            content_generated_url = mitigation["attack_id"].replace(".", "/")
            add_hyperlink(techniqueurl,mitigation['attack_id'],"https://attack.mitre.org/techniques/" + content_generated_url)
            row_cells = table_mitigations.add_row().cells
            row_cells[0].merge(row_cells[2])
            row_cells[0].text = mitigation["description"]
            row_cells = table_mitigations.add_row().cells
            row_cells[0].merge(row_cells[2])
            row_cells[0].text = mitigation["nist_control"]
        document.save(file_docx_mitigations)
    elif switch_control_mapping_selection == "CX":
        document = Document(file_docx_template)
        document.add_heading("Mitigations/Controls",1)
        document.add_paragraph("Mitigations represent security concepts and classes of technologies that can be used to prevent (Sub)-Techniques from being successfully executed.")
        document.add_paragraph()
        document.add_heading("Mitigations Resume",2)
        document.add_paragraph()
        for mitigation in array_obj_sorted_mitigations:
            bulleted = document.add_paragraph(style='List Bullet')
            bulleted.add_run(mitigation["description"])
        document.add_page_break()  
        document.add_heading("Mitigations Overview",2)
        document.add_paragraph()
        document.add_paragraph("The mitigations listed below are mapped with the CIS Controls® v8. This mapping demonstrates which Controls are supported with the implementation of the corresponding Mitigations.")
        table_mitigations = document.add_table(rows=0,cols=3)
        table_mitigations.style = 'Table Grid'
        row_cells = table_mitigations.add_row().cells
        row_cells[0].paragraphs[0].add_run('Mitigation ID: Name').bold = True
        row_cells[1].paragraphs[0].add_run('Mitigation URL').bold = True
        row_cells[2].paragraphs[0].add_run('Covered ATT&CK® Technique').bold = True
        row_cells = table_mitigations.add_row().cells
        row_cells[0].merge(row_cells[2])
        row_cells[0].paragraphs[0].add_run('Description').bold = True
        row_cells = table_mitigations.add_row().cells
        row_cells[0].merge(row_cells[2])
        row_cells[0].paragraphs[0].add_run('CIS Controls® v8').bold = True
        for mitigation in array_obj_sorted_mitigations:
            row_cells = table_mitigations.add_row().cells
            row_cells[0].paragraphs[0].add_run(mitigation["external_id"] + ": " + mitigation["name"]).bold = True
            mitigationurl = row_cells[1].paragraphs[0]
            add_hyperlink(mitigationurl,mitigation['external_id'],mitigation['url'])
            techniqueurl = row_cells[2].paragraphs[0]
            content_generated_url = mitigation["attack_id"].replace(".", "/")
            add_hyperlink(techniqueurl,mitigation['attack_id'],"https://attack.mitre.org/techniques/" + content_generated_url)
            row_cells = table_mitigations.add_row().cells
            row_cells[0].merge(row_cells[2])
            row_cells[0].text = mitigation["description"]
            row_cells = table_mitigations.add_row().cells
            row_cells[0].merge(row_cells[2])
            row_cells[0].text = mitigation["cis_control"]
        document.add_page_break()  
        document.add_heading("CIS Controls® Implementation Priority Guideline",2)
        document.add_paragraph("Below list presents a possible implementation priority, based on the lowest implementation groups where the CIS Control® is associated with and the weight of that specific CIS Control® in the mapping with the identified ATT&CK® (Sub-)Techniques and their associated Mitigations.")
        table_cis_controls_prio = document.add_table(rows=0,cols=4)
        table_cis_controls_prio.style = 'Table Grid'
        row_cells = table_cis_controls_prio.add_row().cells
        row_cells[0].paragraphs[0].add_run('Control® ID').bold = True
        row_cells[1].paragraphs[0].add_run('Control® Description').bold = True
        row_cells[2].paragraphs[0].add_run('IG').bold = True
        row_cells[3].paragraphs[0].add_run('Relative Weight').bold = True
        for mitigation in array_obj_complete_cis_controls_prio_sorted:
            row_cells = table_cis_controls_prio.add_row().cells
            row_cells[0].paragraphs[0].add_run(mitigation["cis_control_id"]).bold = True
            row_cells[1].text = mitigation["cis_control_name"]
            row_cells[2].text = mitigation["cis_control_ig"]
            row_cells[3].text = (str(mitigation["cis_control_count"]))
        document.save(file_docx_mitigations)
    elif switch_control_mapping_selection == "XX":
        document = Document(file_docx_template)
        document.add_heading("Mitigations/Controls",1)
        document.add_paragraph("Mitigations represent security concepts and classes of technologies that can be used to prevent (Sub)-Techniques from being successfully executed.")
        document.add_paragraph()
        document.add_heading("Mitigations Resume",2)
        document.add_paragraph()
        for mitigation in array_obj_sorted_mitigations:
            bulleted = document.add_paragraph(style='List Bullet')
            bulleted.add_run(mitigation["description"])
        document.add_page_break()  
        document.add_heading("Mitigations Overview",2)
        document.add_paragraph()
        table_mitigations = document.add_table(rows=0,cols=3)
        table_mitigations.style = 'Table Grid'
        row_cells = table_mitigations.add_row().cells
        row_cells[0].paragraphs[0].add_run('Mitigation ID: Name').bold = True
        row_cells[1].paragraphs[0].add_run('Mitigation URL').bold = True
        row_cells[2].paragraphs[0].add_run('Covered ATT&CK® Technique').bold = True
        row_cells = table_mitigations.add_row().cells
        row_cells[0].merge(row_cells[2])
        row_cells[0].paragraphs[0].add_run('Description').bold = True
        for mitigation in array_obj_sorted_mitigations:
            row_cells = table_mitigations.add_row().cells
            row_cells[0].paragraphs[0].add_run(mitigation["external_id"] + ": " + mitigation["name"]).bold = True
            mitigationurl = row_cells[1].paragraphs[0]
            add_hyperlink(mitigationurl,mitigation['external_id'],mitigation['url'])
            techniqueurl = row_cells[2].paragraphs[0]
            content_generated_url = mitigation["attack_id"].replace(".", "/")
            add_hyperlink(techniqueurl,mitigation['attack_id'],"https://attack.mitre.org/techniques/" + content_generated_url)
            row_cells = table_mitigations.add_row().cells
            row_cells[0].merge(row_cells[2])
            row_cells[0].text = mitigation["description"]
        document.save(file_docx_mitigations)
    else:
        pass

def new_attackdetectionsconstruct():
    array_obj_complete_detections = [obj for obj in array_obj_complete_attack["objects"] if obj.get("x_mitre_deprecated") != True and obj.get("revoked") != True and obj.get("relationship_type") == "detects"]
    array_obj_filtered_detections = [obj for obj in array_obj_complete_detections if obj["target_ref"] in [attack_pattern["id"] for attack_pattern in array_obj_filtered_mapping_external_id_attack_pattern]]
    array_obj_filtered_mitigations_detections = []
    for detection in array_obj_filtered_detections:
        obj_detection_property_guid = detection["source_ref"]
        obj_detection_property_description = re.sub(r'\(Citation:.*\)', '', detection["description"])
        obj_detection_property_description = re.sub(r'<h4>\s+', '<h4>', obj_detection_property_description)
        obj_detection_property_description = re.sub(r'<h5>\s+', '<h5>', obj_detection_property_description)
        obj_detection_property_description = re.sub(r'\s*</h4>', '</h4>', obj_detection_property_description)
        obj_detection_property_description = re.sub(r'\s*</h5>', '</h5>', obj_detection_property_description)
        description_pattern = re.compile(r'^(.*?)\n\n<h4>', re.DOTALL)
        implementation_pattern = re.compile(r'<h4>Implementation\s*\d*\s*:\s*(.*?)</h4>', re.DOTALL)
        pseudocode_pattern = re.compile(r'<h5>Detection Pseudocode</h5>\n<code>(.*?)</code>', re.DOTALL)
    #    notes_pattern = re.compile(r'<h4>Detection Notes<\/h4>\n\n(.*?)\n', re.DOTALL)
        description_short = description_pattern.findall(obj_detection_property_description)
        if len(description_short) == 0:
            description_short.append(obj_detection_property_description)
        description_implementations = implementation_pattern.findall(obj_detection_property_description)
        description_pseudocodes = pseudocode_pattern.findall(obj_detection_property_description)
    #    description_notes = notes_pattern.findall(obj_detection_property_description)
        description_detections = []
        for i in range(len(description_implementations)):
            description_detection = {
                "implementation": description_implementations[i].strip(),
                "pseudocode": description_pseudocodes[i].strip()
    #            "notes": description_notes[i].strip()
            }
            description_detections.append(description_detection)
        attack_detection_attack_pattern = next((attack_pattern for attack_pattern in array_obj_filtered_mapping_external_id_attack_pattern if attack_pattern["id"] == detection["target_ref"]), None)
        detection_component_block = next((obj for obj in array_obj_complete_attack["objects"] if obj.get("type") == "x-mitre-data-component" and obj.get("id") == obj_detection_property_guid), None)
        detection_data_source = detection_component_block.get("x_mitre_data_source_ref")
        detection_data_source_block = next((obj for obj in array_obj_complete_attack["objects"] if obj.get("type") == "x-mitre-data-source" and obj.get("id") == detection_data_source), None)
        detection_data_source_block_id = next((ref for ref in detection_data_source_block.get("external_references", []) if ref.get("source_name") == "mitre-attack"), None)
        array_row = {
            "name": detection_component_block.get("name"),
            "external_id": detection_data_source_block_id.get("external_id"),
            "url": detection_data_source_block_id.get("url").replace("-", ""),
            "description": obj_detection_property_description,
            "reduced_description": description_short,
            "car_pseudocode": description_detections,
            "platforms": detection_data_source_block.get("x_mitre_platforms"),
            "collection_layers": detection_data_source_block.get("x_mitre_collection_layers"),
            "attack_id": attack_detection_attack_pattern.get("external_id")
        }
        array_obj_filtered_mitigations_detections.append(array_row)
    array_obj_sorted_detections = sorted(array_obj_filtered_mitigations_detections, key=lambda x: (x["external_id"], x["name"], x["attack_id"]), reverse=False)
    grouped_detections = {}
    for detection in array_obj_sorted_detections:
        key = (detection["external_id"], detection["name"])
        if key not in grouped_detections:
            grouped_detections[key] = []
        grouped_detections[key].append(detection)

    array_obj_condensed_detections = []
    for key, group in grouped_detections.items():
        condensed_detection = {
            "name": group[0]["name"],
            "external_id": group[0]["external_id"],
            "url": group[0]["url"],
            "attack_id": [detection["attack_id"] for detection in group],
            "platforms": group[0]["platforms"],
            "collection_layers": group[0]["collection_layers"],
            "description": [],
            "car_pseudocode": [],
            "combined_attack": group[0]["attack_id"]
        }
        for detection in group:
            condensed_detection["car_pseudocode"].extend(detection["car_pseudocode"])
            condensed_detection["description"].extend(detection["reduced_description"])
        array_obj_condensed_detections.append(condensed_detection)
    globals()["array_obj_condensed_detections"] = array_obj_condensed_detections
    globals()["array_obj_filtered_mitigations_detections"] = array_obj_filtered_mitigations_detections

def new_attackdocdetections():
    new_attackdetectionsconstruct()
    file_docx_detections = os.path.join(case_path, document_prefix + "detections.docx")
    document = Document(file_docx_template)
    document.add_heading("Detections",1)
    document.add_paragraph("Detections are based on data sources and their components associated with the identified (Sub-)Techniques required to create detections where the mitigations/controls prove to be impossible to implement or inadequate.\nThe table includes the mapping with the Open Source Security Events Metadata Detection Model (OSSEM-DM) and extracted information from MITRE Cyber Analytics Repository (CAR) where available. It facilitates the detection of adversary techniques.\nThe provided information may help or drive the development of detection rules for adversary actions mapped to the MITRE ATT&CK knowledge base.")
    document.add_paragraph()
    table = document.add_table(rows=0,cols=3)
    table.style = 'Table Grid'
    row_cells = table.add_row().cells
    row_cells[0].paragraphs[0].add_run('Detection ID: Name').bold = True
    row_cells[1].paragraphs[0].add_run('Detection URL').bold = True
    row_cells[2].paragraphs[0].add_run('Covered ATT&CK® Technique').bold = True
    row_cells = table.add_row().cells
    row_cells[0].paragraphs[0].add_run('Platforms').bold = True
    row_cells[1].merge(row_cells[2])
    row_cells[1].paragraphs[0].add_run('Collection Layers').bold = True
    row_cells = table.add_row().cells
    row_cells[0].merge(row_cells[2])
    row_cells[0].paragraphs[0].add_run('Description').bold = True
    row_cells = table.add_row().cells
    row_cells[0].merge(row_cells[2])
    row_cells = table.add_row().cells
    row_cells[0].merge(row_cells[2])
    row_cells[0].paragraphs[0].add_run('CAR Pseudocode').bold = True
    row_cells = table.add_row().cells
    row_cells[0].merge(row_cells[2])
    row_cells = table.add_row().cells
    row_cells[0].merge(row_cells[2])
    row_cells[0].paragraphs[0].add_run('Source - Relationship - Target').bold = True
    row_cells = table.add_row().cells
    row_cells[0].merge(row_cells[2])
    row_cells[0].paragraphs[0].add_run('Log Source/Channel').bold = True
    row_cells = table.add_row().cells
    row_cells[0].merge(row_cells[2])
    row_cells[0].paragraphs[0].add_run('EventID - Event Name | Defender Advanced Hunting Schema/ActionType filter').bold = True
    row_cells = table.add_row().cells
    row_cells[0].merge(row_cells[2])
    row_cells[0].paragraphs[0].add_run('Platform/Audit Category/Audit Subcategory : Filter').bold = True
    document.add_paragraph()
    for item in array_obj_condensed_detections:
        document.add_page_break()
        table = document.add_table(rows=0,cols=3)
        table.style = 'Table Grid'
        row_cells = table.add_row().cells
        run = row_cells[0].paragraphs[0].add_run(item['external_id'] + ": " + item['name'])
        run.bold = True
        run.font.color.rgb = RGBColor(218,21,114)
        datasourceurl = row_cells[1].paragraphs[0]
        add_hyperlink(datasourceurl,item['external_id'],item['url'])
        row_cells[2].text = ", ".join(item['attack_id'])
        row_cells = table.add_row().cells
        row_cells[0].text = ", ".join(item['platforms'])
        row_cells[1].merge(row_cells[2])
        row_cells[1].text = ", ".join(item['collection_layers'])
        row_cells = table.add_row().cells
        row_cells[0].merge(row_cells[2])
        para = row_cells[0].paragraphs[0]
        description = sorted(set(item['description']))
        text = "\n".join(description)
        process_text_with_links_code(text, para)
        var_car_pseudocode_elements = len(item["car_pseudocode"])
        table_pseudocode = document.add_table(rows=0,cols=1)
        table_pseudocode.style = 'Table Grid'
        table_pseudocode.add_row()
        row_cells = table_pseudocode.add_row().cells
        if var_car_pseudocode_elements == 0:
            run = row_cells[0].paragraphs[0].add_run('No CAR Pseudocode Information available.')
            run.bold = True
            run.font.color.rgb = RGBColor(140,14,74)
        else:
            run = row_cells[0].paragraphs[0].add_run('CAR Pseudocode Information:')
            run.bold = True
            run.italic = True
            run.font.color.rgb = RGBColor(218,21,114)
            for c in (item["car_pseudocode"]):
                table_pseudocode.add_row()
                row_cells = table_pseudocode.add_row().cells
                row_cells[0].paragraphs[0].add_run(c['implementation']).bold = True
                row_cells = table_pseudocode.add_row().cells
                row_cells[0].paragraphs[0].add_run(c['pseudocode']).bold = False
        array_obj_filtered_ossem_data = [ossem_obj for ossem_obj in array_obj_complete_ossem_mapping if ossem_obj["technique_id"] == item["combined_attack"] and str(ossem_obj["data_component"]) == item["name"].lower()]
        var_ossem_elements = len(array_obj_filtered_ossem_data)
        table_ossem = document.add_table(rows=0,cols=1)
        table_ossem.style = 'Table Grid'
        table_ossem.add_row()
        row_cells = table_ossem.add_row().cells
        if var_ossem_elements == 0:
            run = row_cells[0].paragraphs[0].add_run('No OSSEM DM Information available.')
            run.bold = True
            run.font.color.rgb = RGBColor(140,14,74)
        else:
            run = row_cells[0].paragraphs[0].add_run('OSSEM DM Information:')
            run.bold = True
            run.italic = True
            run.font.color.rgb = RGBColor(218,21,114)
            for j in array_obj_filtered_ossem_data:
                table_ossem.add_row()
                row_cells = table_ossem.add_row().cells
                row_cells[0].paragraphs[0].add_run('Source - Relationship - Target: ' + j['name']).bold = True
                row_cells = table_ossem.add_row().cells
                if j['log_source'] == "sysmon" or j['log_source'] == "Microsoft Defender for Endpoint":
                    row_cells[0].text = "Log Source: " + j['log_source']
                elif j['log_source'] == "Microsoft-Windows-Sysmon":
                    row_cells[0].text = "Log Source/Channel: " + (str(j['channel']))
                else:
                    if str(j['channel']) == 'nan':
                        row_cells[0].text = "Log Source/Channel: " + j['log_source']
                    else:
                        row_cells[0].text = "Log Source/Channel: " + j['log_source']  + "/" + (str(j['channel']))
                row_cells = table_ossem.add_row().cells
                if j['log_source'] == "Microsoft Defender for Endpoint":
                    row_cells[0].text = "Defender Advanced Hunting Schema/ActionType filter: " + j['event_id'] + "/" + j['filter_in'][0]['ActionType']
                else:
                    row_cells[0].text = "EventID - Event Name: " + str(j['event_id']) + " - " + j['event_name']
                row_cells = table_ossem.add_row().cells
                if str(j['audit_sub_category']) == 'nan':
                    if str(j['audit_category']) == 'nan':
                        row_cells[0].text = "Platform: " + j['event_platform']
                    else:
                        if str(j['filter_in']) == 'nan':
                            row_cells[0].text = "Platform/Audit Category: " + j['event_platform'] + "/" + j['audit_category']
                        else:
                            row_cells[0].text = "Platform/Audit Category : Filter: " + j['event_platform'] + "/" + j['audit_category'] + " : " + str(j['filter_in'])
                else:
                    if str(j['filter_in']) == 'nan':
                        row_cells[0].text = "Platform/Audit Category/Audit Subcategory: " + j['event_platform'] + "/" + j['audit_category'] + "/" + j['audit_sub_category']
                    else:
                        row_cells[0].text = "Platform/Audit Category/Audit Subcategory : Filter: " + j['event_platform'] + "/" + j['audit_category'] + "/" + j['audit_sub_category'] + " : " + str(j['filter_in'])
    document.save(file_docx_detections)

def new_attackdocvalidations():
    file_docx_validations = os.path.join(case_path, document_prefix + "validations.docx")
    array_obj_complete_validation = []
    array_obj_complete_validation = [technique for technique in array_obj_complete_atomicred_mapping["techniques"] if technique["techniqueID"] in [attack["attack_id"] for attack in list_obj_selected_attack_techniques]]
    document = Document(file_docx_template)
    document.add_heading("Validations",1)
    document.add_paragraph("Validations are based on Atomic Red Team tests. The references point to the available tests for the given Techniques. These are not to be considered as providing a complete coverage of all possible ways to simulate the effects of a given Technique. It facilitates validation your mitigations and detections for your environment.")
    document.add_paragraph()
    table = document.add_table(rows=0,cols=2)
    table.style = 'Table Grid'
    row_cells = table.add_row().cells
    row_cells[0].paragraphs[0].add_run('Atomic Red Team test URL').bold = True
    row_cells[1].paragraphs[0].add_run('Score').bold = True
    for item in array_obj_complete_validation:
            row_cells = table.add_row().cells
            validationsourceeurl = row_cells[0].paragraphs[0]
            add_hyperlink(validationsourceeurl,"Atomic Red Team test for " + item['techniqueID'],item['links'][0]['url'])
            row_cells[1].text = str(item['score'])
    document.save(file_docx_validations)

def new_ctidattackflow(ctid_assets=None):
    file_afb_ctid_flow = os.path.join(case_path, document_prefix + "ctid_attack_flow.afb")
    var_obj_flow_property_GUID = str(uuid.uuid4())
    var_obj_flow_property_background_colour = "#ffffff"
    var_obj_flow_property_grid_colour = "#f5f5f5"
    var_obj_flow_objects_property_anchor_markers_colour = "#fb6fa5"
    var_obj_flow_objects_property_anchor_hover_colour = "rgba(200, 88, 135, 0.25)"
    var_obj_flow_objects_property_box_colour = "#fefefe"
    now = datetime.utcnow()
    current_time = now.strftime('%Y-%m-%dT%H:%M:%S') + 'Z'
    obj_flow_template_header_0  = '''{"version":"2.0.1","id":"'''
    obj_flow_template_header_1  = var_obj_flow_property_GUID
    obj_flow_template_header_2  = '''","schema":{"page_template":"flow","templates":[{"id":"@__builtin__page","type":7,"role":0,"grid":[10,10],"properties":{"name":{"type":2,"value":"Untitled Document","is_primary":true}},"style":{"grid_color":"'''
    obj_flow_template_header_3 = var_obj_flow_property_grid_colour
    obj_flow_template_header_4 = '''","background_color":"'''
    obj_flow_template_header_5 = var_obj_flow_property_background_colour
    obj_flow_template_header_6 = '''","drop_shadow":{"color":"rgba(0,0,0,.4)","offset":[3,3]}}},{"id":"@__builtin__anchor","type":0,"role":0,"radius":10,"line_templates":{"0":"@__builtin__line_horizontal_elbow","1":"@__builtin__line_vertical_elbow"},"style":{"color":"'''
    obj_flow_template_header_7 = var_obj_flow_objects_property_anchor_hover_colour
    obj_flow_template_header_8 = '''"}},{"id":"@__builtin__line_handle","type":4,"role":0,"style":{"radius":6,"fill_color":"#fedb22","stroke_color":"#fefefe","stroke_width":1.5}},{"id":"@__builtin__line_source","type":3,"role":12288,"style":{"radius":6,"fill_color":"#fedb22","stroke_color":"#141414","stroke_width":1.5}},{"id":"@__builtin__line_target","type":3,"role":16384,"style":{"radius":6,"fill_color":"#fedb22","stroke_color":"#141414","stroke_width":1.5}},{"id":"@__builtin__line_horizontal_elbow","namespace":"horizontal_elbow","type":5,"role":8192,"hitbox_width":20,"line_handle_template":"@__builtin__line_handle","line_ending_template":{"source":"@__builtin__line_source","target":"@__builtin__line_target"},"style":{"width":5,"cap_size":16,"color":"#646464","select_color":"#646464"}},{"id":"@__builtin__line_vertical_elbow","namespace":"vertical_elbow","type":6,"role":8192,"hitbox_width":20,"line_handle_template":"@__builtin__line_handle","line_ending_template":{"source":"@__builtin__line_source","target":"@__builtin__line_target"},"style":{"width":5,"cap_size":16,"color":"#646464","select_color":"#646464"}},{"id":"flow","type":7,"role":4096,"grid":[10,10],"properties":{"name":{"type":2,"value":"Untitled Document","is_primary":true},"description":{"type":2},"author":{"type":6,"form":{"name":{"type":2,"is_primary":true},"identity_class":{"type":4,"options":{"type":5,"form":{"type":2},"value":[["individual","Individual"],["group","Group"],["system","System"],["organization","Organization"],["class","Class"],["unknown","Unknown"]]}},"contact_information":{"type":2}}},"scope":{"type":4,"options":{"type":5,"form":{"type":2},"value":[["incident","Incident"],["campaign","Campaign"],["threat-actor","Threat Actor"],["malware","Malware"],["other","Other"]]},"value":"incident"},"external_references":{"type":5,"form":{"type":6,"form":{"source_name":{"type":2,"is_primary":true,"is_required":true},"description":{"type":2},"url":{"type":2}}}},"created":{"type":3,"value":"'''
    obj_flow_template_header_9 = '''current_time","is_visible":false}},"style":{"grid_color":"'''
    obj_flow_template_header_10 = var_obj_flow_property_grid_colour
    obj_flow_template_header_11 = '''","background_color":"'''
    obj_flow_template_header_12 = var_obj_flow_property_background_colour
    obj_flow_template_header_13 = '''","drop_shadow":{"color":"rgba(0,0,0,.4)","offset":[3,3]}}},{"id":"true_anchor","type":0,"role":0,"radius":10,"line_templates":{"0":"@__builtin__line_horizontal_elbow","1":"@__builtin__line_vertical_elbow"},"style":{"color":"rgba(255, 255, 255, 0.25)"}},{"id":"false_anchor","type":0,"role":0,"radius":10,"line_templates":{"0":"@__builtin__line_horizontal_elbow","1":"@__builtin__line_vertical_elbow"},"style":{"color":"rgba(255, 255, 255, 0.25)"}},{"id":"action","namespace":"attack_flow.action","type":2,"role":4096,"properties":{"name":{"type":2,"is_primary":true,"is_required":true},"tactic_id":{"type":2},"tactic_ref":{"type":2},"technique_id":{"type":2},"technique_ref":{"type":2},"description":{"type":2},"confidence":{"type":4,"options":{"type":5,"form":{"type":6,"form":{"text":{"type":2,"is_primary":true},"value":{"type":0}}},"value":[["speculative",{"text":"Speculative","value":0}],["very-doubtful",{"text":"Very Doubtful","value":10}],["doubtful",{"text":"Doubtful","value":30}],["even-odds",{"text":"Even Odds","value":50}],["probable",{"text":"Probable","value":70}],["very-probable",{"text":"Very Probable","value":90}],["certain",{"text":"Certain","value":100}]]},"value":null},"execution_start":{"type":3},"execution_end":{"type":3}},"anchor_template":"@__builtin__anchor","style":{"max_width":320,"head":{"fill_color":"#637bc9","stroke_color":"#708ce6","one_title":{"title":{"font":{"family":"Inter","size":"10.5pt","weight":800},"color":"#d8d8d8"}},"two_title":{"title":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#d8d8d8","padding":8},"subtitle":{"font":{"family":"Inter","size":"13pt","weight":800},"color":"#d8d8d8","line_height":23}},"vertical_padding":14},"body":{"fill_color":"'''
    obj_flow_template_header_14 = var_obj_flow_objects_property_box_colour
    obj_flow_template_header_15 = '''","stroke_color":"#383838","field_name":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#b3b3b3","padding":12},"field_value":{"font":{"family":"Inter","size":"10.5pt"},"color":"#bfbfbf","line_height":20,"padding":22},"vertical_padding":18},"select_outline":{"color":"#e6d845","padding":4,"border_radius":9},"anchor_markers":{"color":"'''
    obj_flow_template_header_16 = var_obj_flow_objects_property_anchor_markers_colour
    obj_flow_template_header_17 = '''","size":3},"border_radius":5,"horizontal_padding":20}},{"id":"asset","namespace":"attack_flow.asset","type":2,"role":4096,"properties":{"name":{"type":2,"is_primary":true,"is_required":true},"description":{"type":2}},"anchor_template":"@__builtin__anchor","style":{"max_width":320,"head":{"fill_color":"#c26130","stroke_color":"#e57339","one_title":{"title":{"font":{"family":"Inter","size":"10.5pt","weight":800},"color":"#d8d8d8"}},"two_title":{"title":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#d8d8d8","padding":8},"subtitle":{"font":{"family":"Inter","size":"13pt","weight":800},"color":"#d8d8d8","line_height":23}},"vertical_padding":14},"body":{"fill_color":"'''
    obj_flow_template_header_18 = var_obj_flow_objects_property_box_colour
    obj_flow_template_header_19 = '''","stroke_color":"#383838","field_name":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#b3b3b3","padding":12},"field_value":{"font":{"family":"Inter","size":"10.5pt"},"color":"#bfbfbf","line_height":20,"padding":22},"vertical_padding":18},"select_outline":{"color":"#e6d845","padding":4,"border_radius":9},"anchor_markers":{"color":"'''
    obj_flow_template_header_20 = var_obj_flow_objects_property_anchor_markers_colour
    obj_flow_template_header_21 = '''","size":3},"border_radius":5,"horizontal_padding":20}},{"id":"condition","namespace":"attack_flow.condition","type":1,"role":4096,"properties":{"description":{"type":2,"is_primary":true,"is_required":true},"pattern":{"type":2},"pattern_type":{"type":2},"pattern_version":{"type":2},"date":{"type":3}},"branches":[{"text":"True","anchor_template":"true_anchor"},{"text":"False","anchor_template":"false_anchor"}],"anchor_template":"@__builtin__anchor","style":{"max_width":320,"head":{"fill_color":"#2a9642","stroke_color":"#32b34e","one_title":{"title":{"font":{"family":"Inter","size":"10.5pt","weight":800},"color":"#d8d8d8"}},"two_title":{"title":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#d8d8d8","padding":8},"subtitle":{"font":{"family":"Inter","size":"13pt","weight":800},"color":"#d8d8d8","line_height":23}},"vertical_padding":14},"body":{"fill_color":"'''
    obj_flow_template_header_22 = var_obj_flow_objects_property_box_colour
    obj_flow_template_header_23 = '''","stroke_color":"#383838","field_name":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#b3b3b3","padding":12},"field_value":{"font":{"family":"Inter","size":"10.5pt"},"color":"#bfbfbf","line_height":20,"padding":22},"vertical_padding":18},"branch":{"font":{"family":"Inter","size":"10.5pt"},"color":"#bfbfbf","vertical_padding":12,"horizontal_padding":30},"select_outline":{"color":"#e6d845","padding":4,"border_radius":9},"anchor_markers":{"color":"'''
    obj_flow_template_header_24 = var_obj_flow_objects_property_anchor_markers_colour
    obj_flow_template_header_25 = '''","size":3},"border_radius":5,"horizontal_padding":20}},{"id":"or","namespace":"attack_flow.OR_operator","type":8,"role":4096,"properties":{"operator":{"type":2,"value":"OR","is_primary":true,"is_visible":false,"is_editable":false}},"anchor_template":"@__builtin__anchor","style":{"max_width":320,"fill_color":"#c94040","stroke_color":"#dd5050","text":{"font":{"family":"Inter","size":"14pt","weight":800},"color":"#d8d8d8","line_height":24},"border_radius":13,"select_outline":{"color":"#e6d845","padding":4,"border_radius":19},"anchor_markers":{"color":"'''
    obj_flow_template_header_26 = var_obj_flow_objects_property_anchor_markers_colour
    obj_flow_template_header_27 = '''","size":3},"vertical_padding":18,"horizontal_padding":35}},{"id":"and","namespace":"attack_flow.AND_operator","type":8,"role":4096,"properties":{"operator":{"type":2,"value":"AND","is_primary":true,"is_visible":false,"is_editable":false}},"anchor_template":"@__builtin__anchor","style":{"max_width":320,"fill_color":"#c94040","stroke_color":"#dd5050","text":{"font":{"family":"Inter","size":"14pt","weight":800},"color":"#d8d8d8","line_height":24},"border_radius":13,"select_outline":{"color":"#e6d845","padding":4,"border_radius":19},"anchor_markers":{"color":"'''
    obj_flow_template_header_28 = var_obj_flow_objects_property_anchor_markers_colour
    obj_flow_template_header_29 = '''","size":3},"vertical_padding":18,"horizontal_padding":35}},{"id":"attack_pattern","namespace":"stix_object.attack_pattern","type":2,"role":4096,"properties":{"name":{"type":2,"is_primary":true,"is_required":true},"description":{"type":2},"aliases":{"type":5,"form":{"type":2}},"kill_chain_phases":{"type":5,"form":{"type":2}}},"anchor_template":"@__builtin__anchor","style":{"max_width":320,"head":{"fill_color":"#737373","stroke_color":"#8c8c8c","one_title":{"title":{"font":{"family":"Inter","size":"10.5pt","weight":800},"color":"#d8d8d8"}},"two_title":{"title":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#d8d8d8","padding":8},"subtitle":{"font":{"family":"Inter","size":"13pt","weight":800},"color":"#d8d8d8","line_height":23}},"vertical_padding":14},"body":{"fill_color":"'''
    obj_flow_template_header_30 = var_obj_flow_objects_property_box_colour
    obj_flow_template_header_31 = '''","stroke_color":"#383838","field_name":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#b3b3b3","padding":12},"field_value":{"font":{"family":"Inter","size":"10.5pt"},"color":"#bfbfbf","line_height":20,"padding":22},"vertical_padding":18},"select_outline":{"color":"#e6d845","padding":4,"border_radius":9},"anchor_markers":{"color":"'''
    obj_flow_template_header_32 = var_obj_flow_objects_property_anchor_markers_colour
    obj_flow_template_header_33 = '''","size":3},"border_radius":5,"horizontal_padding":20}},{"id":"campaign","namespace":"stix_object.campaign","type":2,"role":4096,"properties":{"name":{"type":2,"is_primary":true,"is_required":true},"description":{"type":2},"aliases":{"type":5,"form":{"type":2}},"first_seen":{"type":3},"last_seen":{"type":3},"objective":{"type":2}},"anchor_template":"@__builtin__anchor","style":{"max_width":320,"head":{"fill_color":"#737373","stroke_color":"#8c8c8c","one_title":{"title":{"font":{"family":"Inter","size":"10.5pt","weight":800},"color":"#d8d8d8"}},"two_title":{"title":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#d8d8d8","padding":8},"subtitle":{"font":{"family":"Inter","size":"13pt","weight":800},"color":"#d8d8d8","line_height":23}},"vertical_padding":14},"body":{"fill_color":"'''
    obj_flow_template_header_34 = var_obj_flow_objects_property_box_colour
    obj_flow_template_header_35 = '''","stroke_color":"#383838","field_name":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#b3b3b3","padding":12},"field_value":{"font":{"family":"Inter","size":"10.5pt"},"color":"#bfbfbf","line_height":20,"padding":22},"vertical_padding":18},"select_outline":{"color":"#e6d845","padding":4,"border_radius":9},"anchor_markers":{"color":"'''
    obj_flow_template_header_36 = var_obj_flow_objects_property_anchor_markers_colour
    obj_flow_template_header_37 = '''","size":3},"border_radius":5,"horizontal_padding":20}},{"id":"course_of_action","namespace":"stix_object.course_of_action","type":2,"role":4096,"properties":{"name":{"type":2,"is_primary":true,"is_required":true},"description":{"type":2},"action_type":{"type":2},"os_execution_envs":{"type":5,"form":{"type":2}},"action_bin":{"type":2}},"anchor_template":"@__builtin__anchor","style":{"max_width":320,"head":{"fill_color":"#737373","stroke_color":"#8c8c8c","one_title":{"title":{"font":{"family":"Inter","size":"10.5pt","weight":800},"color":"#d8d8d8"}},"two_title":{"title":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#d8d8d8","padding":8},"subtitle":{"font":{"family":"Inter","size":"13pt","weight":800},"color":"#d8d8d8","line_height":23}},"vertical_padding":14},"body":{"fill_color":"'''
    obj_flow_template_header_38 = var_obj_flow_objects_property_box_colour
    obj_flow_template_header_39 = '''","stroke_color":"#383838","field_name":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#b3b3b3","padding":12},"field_value":{"font":{"family":"Inter","size":"10.5pt"},"color":"#bfbfbf","line_height":20,"padding":22},"vertical_padding":18},"select_outline":{"color":"#e6d845","padding":4,"border_radius":9},"anchor_markers":{"color":"'''
    obj_flow_template_header_40 = var_obj_flow_objects_property_anchor_markers_colour
    obj_flow_template_header_41 = '''","size":3},"border_radius":5,"horizontal_padding":20}},{"id":"grouping","namespace":"stix_object.grouping","type":2,"role":4096,"properties":{"name":{"type":2,"is_primary":true},"description":{"type":2},"context":{"type":2,"is_required":true}},"anchor_template":"@__builtin__anchor","style":{"max_width":320,"head":{"fill_color":"#737373","stroke_color":"#8c8c8c","one_title":{"title":{"font":{"family":"Inter","size":"10.5pt","weight":800},"color":"#d8d8d8"}},"two_title":{"title":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#d8d8d8","padding":8},"subtitle":{"font":{"family":"Inter","size":"13pt","weight":800},"color":"#d8d8d8","line_height":23}},"vertical_padding":14},"body":{"fill_color":"'''
    obj_flow_template_header_42 = var_obj_flow_objects_property_box_colour
    obj_flow_template_header_43 = '''","stroke_color":"#383838","field_name":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#b3b3b3","padding":12},"field_value":{"font":{"family":"Inter","size":"10.5pt"},"color":"#bfbfbf","line_height":20,"padding":22},"vertical_padding":18},"select_outline":{"color":"#e6d845","padding":4,"border_radius":9},"anchor_markers":{"color":"'''
    obj_flow_template_header_44 = var_obj_flow_objects_property_anchor_markers_colour
    obj_flow_template_header_45 = '''","size":3},"border_radius":5,"horizontal_padding":20}},{"id":"identity","namespace":"stix_object.identity","type":2,"role":4096,"properties":{"name":{"type":2,"is_required":true,"is_primary":true},"description":{"type":2},"roles":{"type":5,"form":{"type":2}},"identity_class":{"type":2,"is_required":true},"sectors":{"type":5,"form":{"type":2}},"contact_information":{"type":2}},"anchor_template":"@__builtin__anchor","style":{"max_width":320,"head":{"fill_color":"#737373","stroke_color":"#8c8c8c","one_title":{"title":{"font":{"family":"Inter","size":"10.5pt","weight":800},"color":"#d8d8d8"}},"two_title":{"title":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#d8d8d8","padding":8},"subtitle":{"font":{"family":"Inter","size":"13pt","weight":800},"color":"#d8d8d8","line_height":23}},"vertical_padding":14},"body":{"fill_color":"'''
    obj_flow_template_header_46 = var_obj_flow_objects_property_box_colour
    obj_flow_template_header_47 = '''","stroke_color":"#383838","field_name":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#b3b3b3","padding":12},"field_value":{"font":{"family":"Inter","size":"10.5pt"},"color":"#bfbfbf","line_height":20,"padding":22},"vertical_padding":18},"select_outline":{"color":"#e6d845","padding":4,"border_radius":9},"anchor_markers":{"color":"'''
    obj_flow_template_header_48 = var_obj_flow_objects_property_anchor_markers_colour
    obj_flow_template_header_49 = '''","size":3},"border_radius":5,"horizontal_padding":20}},{"id":"indicator","namespace":"stix_object.indicator","type":2,"role":4096,"properties":{"name":{"type":2,"is_primary":true},"description":{"type":2},"indicator_types":{"type":5,"form":{"type":2,"is_required":true}},"pattern":{"type":2,"is_required":true},"pattern_type":{"type":2,"is_required":true},"patter_version":{"type":2},"valid_from":{"type":3,"is_required":true},"valid_until":{"type":3},"kill_chain_phases":{"type":5,"form":{"type":2}}},"anchor_template":"@__builtin__anchor","style":{"max_width":320,"head":{"fill_color":"#737373","stroke_color":"#8c8c8c","one_title":{"title":{"font":{"family":"Inter","size":"10.5pt","weight":800},"color":"#d8d8d8"}},"two_title":{"title":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#d8d8d8","padding":8},"subtitle":{"font":{"family":"Inter","size":"13pt","weight":800},"color":"#d8d8d8","line_height":23}},"vertical_padding":14},"body":{"fill_color":"'''
    obj_flow_template_header_50 = var_obj_flow_objects_property_box_colour
    obj_flow_template_header_51 = '''","stroke_color":"#383838","field_name":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#b3b3b3","padding":12},"field_value":{"font":{"family":"Inter","size":"10.5pt"},"color":"#bfbfbf","line_height":20,"padding":22},"vertical_padding":18},"select_outline":{"color":"#e6d845","padding":4,"border_radius":9},"anchor_markers":{"color":"'''
    obj_flow_template_header_52 = var_obj_flow_objects_property_anchor_markers_colour
    obj_flow_template_header_53 = '''","size":3},"border_radius":5,"horizontal_padding":20}},{"id":"infrastructure","namespace":"stix_object.infrastructure","type":2,"role":4096,"properties":{"name":{"type":2,"is_primary":true,"is_required":true},"description":{"type":2},"infrastructure_types":{"type":5,"form":{"type":2,"is_required":true}},"aliases":{"type":5,"form":{"type":2}},"kill_chain_phases":{"type":5,"form":{"type":2}},"first_seen":{"type":3},"last_seen":{"type":3}},"anchor_template":"@__builtin__anchor","style":{"max_width":320,"head":{"fill_color":"#737373","stroke_color":"#8c8c8c","one_title":{"title":{"font":{"family":"Inter","size":"10.5pt","weight":800},"color":"#d8d8d8"}},"two_title":{"title":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#d8d8d8","padding":8},"subtitle":{"font":{"family":"Inter","size":"13pt","weight":800},"color":"#d8d8d8","line_height":23}},"vertical_padding":14},"body":{"fill_color":"'''
    obj_flow_template_header_54 = var_obj_flow_objects_property_box_colour
    obj_flow_template_header_55 = '''","stroke_color":"#383838","field_name":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#b3b3b3","padding":12},"field_value":{"font":{"family":"Inter","size":"10.5pt"},"color":"#bfbfbf","line_height":20,"padding":22},"vertical_padding":18},"select_outline":{"color":"#e6d845","padding":4,"border_radius":9},"anchor_markers":{"color":"'''
    obj_flow_template_header_56 = var_obj_flow_objects_property_anchor_markers_colour
    obj_flow_template_header_57 = '''","size":3},"border_radius":5,"horizontal_padding":20}},{"id":"intrusion_set","namespace":"stix_object.intrusion_set","type":2,"role":4096,"properties":{"name":{"type":2,"is_primary":true,"is_required":true},"description":{"type":2},"aliases":{"type":5,"form":{"type":2,"is_required":true}},"first_seen":{"type":3},"last_seen":{"type":3},"goals":{"type":5,"form":{"type":2}},"resource_level":{"type":2},"primary_motivation":{"type":2},"secondary_motivations":{"type":5,"form":{"type":2}}},"anchor_template":"@__builtin__anchor","style":{"max_width":320,"head":{"fill_color":"#737373","stroke_color":"#8c8c8c","one_title":{"title":{"font":{"family":"Inter","size":"10.5pt","weight":800},"color":"#d8d8d8"}},"two_title":{"title":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#d8d8d8","padding":8},"subtitle":{"font":{"family":"Inter","size":"13pt","weight":800},"color":"#d8d8d8","line_height":23}},"vertical_padding":14},"body":{"fill_color":"'''
    obj_flow_template_header_58 = var_obj_flow_objects_property_box_colour
    obj_flow_template_header_59 = '''","stroke_color":"#383838","field_name":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#b3b3b3","padding":12},"field_value":{"font":{"family":"Inter","size":"10.5pt"},"color":"#bfbfbf","line_height":20,"padding":22},"vertical_padding":18},"select_outline":{"color":"#e6d845","padding":4,"border_radius":9},"anchor_markers":{"color":"'''
    obj_flow_template_header_60 = var_obj_flow_objects_property_anchor_markers_colour
    obj_flow_template_header_61 = '''","size":3},"border_radius":5,"horizontal_padding":20}},{"id":"location","namespace":"stix_object.location","type":2,"role":4096,"properties":{"name":{"type":2,"is_primary":true},"description":{"type":2},"latitude":{"type":1,"min":-90,"max":90},"longitude":{"type":1,"min":-180,"max":180},"precision":{"type":1},"region":{"type":2},"country":{"type":2},"administrative_area":{"type":2},"city":{"type":2},"street_address":{"type":2},"postal_code":{"type":2}},"anchor_template":"@__builtin__anchor","style":{"max_width":320,"head":{"fill_color":"#737373","stroke_color":"#8c8c8c","one_title":{"title":{"font":{"family":"Inter","size":"10.5pt","weight":800},"color":"#d8d8d8"}},"two_title":{"title":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#d8d8d8","padding":8},"subtitle":{"font":{"family":"Inter","size":"13pt","weight":800},"color":"#d8d8d8","line_height":23}},"vertical_padding":14},"body":{"fill_color":"'''
    obj_flow_template_header_62 = var_obj_flow_objects_property_box_colour
    obj_flow_template_header_63 = '''","stroke_color":"#383838","field_name":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#b3b3b3","padding":12},"field_value":{"font":{"family":"Inter","size":"10.5pt"},"color":"#bfbfbf","line_height":20,"padding":22},"vertical_padding":18},"select_outline":{"color":"#e6d845","padding":4,"border_radius":9},"anchor_markers":{"color":"'''
    obj_flow_template_header_64 = var_obj_flow_objects_property_anchor_markers_colour
    obj_flow_template_header_65 = '''","size":3},"border_radius":5,"horizontal_padding":20}},{"id":"malware","namespace":"stix_object.malware","type":2,"role":4096,"properties":{"name":{"type":2,"is_primary":true},"description":{"type":2},"malware_types":{"type":5,"form":{"type":2,"is_required":true}},"is_family":{"type":4,"options":{"type":5,"form":{"type":2},"value":[["true","True"],["false","False"]]},"is_required":true},"aliases":{"type":5,"form":{"type":2}},"kill_chain_phases":{"type":5,"form":{"type":2}},"first_seen":{"type":3},"last_seen":{"type":3},"os_execution_envs":{"type":5,"form":{"type":2}},"architecture_execution_envs":{"type":5,"form":{"type":2}},"implementation_languages":{"type":5,"form":{"type":2}},"capabilities":{"type":5,"form":{"type":2}}},"anchor_template":"@__builtin__anchor","style":{"max_width":320,"head":{"fill_color":"#737373","stroke_color":"#8c8c8c","one_title":{"title":{"font":{"family":"Inter","size":"10.5pt","weight":800},"color":"#d8d8d8"}},"two_title":{"title":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#d8d8d8","padding":8},"subtitle":{"font":{"family":"Inter","size":"13pt","weight":800},"color":"#d8d8d8","line_height":23}},"vertical_padding":14},"body":{"fill_color":"'''
    obj_flow_template_header_66 = var_obj_flow_objects_property_box_colour
    obj_flow_template_header_67 = '''","stroke_color":"#383838","field_name":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#b3b3b3","padding":12},"field_value":{"font":{"family":"Inter","size":"10.5pt"},"color":"#bfbfbf","line_height":20,"padding":22},"vertical_padding":18},"select_outline":{"color":"#e6d845","padding":4,"border_radius":9},"anchor_markers":{"color":"'''
    obj_flow_template_header_68 = var_obj_flow_objects_property_anchor_markers_colour
    obj_flow_template_header_69 = '''","size":3},"border_radius":5,"horizontal_padding":20}},{"id":"malware_analysis","namespace":"stix_object.malware_analysis","type":2,"role":4096,"properties":{"product":{"type":2,"is_primary":true,"is_required":true},"version":{"type":2},"configuration_version":{"type":2},"modules":{"type":5,"form":{"type":2}},"analysis_engine_version":{"type":2},"analysis_definition_version":{"type":2},"submitted":{"type":3},"analysis_started":{"type":3},"analysis_ended":{"type":3},"av_result":{"type":2}},"anchor_template":"@__builtin__anchor","style":{"max_width":320,"head":{"fill_color":"#737373","stroke_color":"#8c8c8c","one_title":{"title":{"font":{"family":"Inter","size":"10.5pt","weight":800},"color":"#d8d8d8"}},"two_title":{"title":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#d8d8d8","padding":8},"subtitle":{"font":{"family":"Inter","size":"13pt","weight":800},"color":"#d8d8d8","line_height":23}},"vertical_padding":14},"body":{"fill_color":"'''
    obj_flow_template_header_70 = var_obj_flow_objects_property_box_colour
    obj_flow_template_header_71 = '''","stroke_color":"#383838","field_name":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#b3b3b3","padding":12},"field_value":{"font":{"family":"Inter","size":"10.5pt"},"color":"#bfbfbf","line_height":20,"padding":22},"vertical_padding":18},"select_outline":{"color":"#e6d845","padding":4,"border_radius":9},"anchor_markers":{"color":"'''
    obj_flow_template_header_72 = var_obj_flow_objects_property_anchor_markers_colour
    obj_flow_template_header_73 = '''","size":3},"border_radius":5,"horizontal_padding":20}},{"id":"note","namespace":"stix_object.note","type":2,"role":4096,"properties":{"abstract":{"type":2,"is_primary":true},"content":{"type":2,"is_required":true},"authors":{"type":5,"form":{"type":2}}},"anchor_template":"@__builtin__anchor","style":{"max_width":320,"head":{"fill_color":"#737373","stroke_color":"#8c8c8c","one_title":{"title":{"font":{"family":"Inter","size":"10.5pt","weight":800},"color":"#d8d8d8"}},"two_title":{"title":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#d8d8d8","padding":8},"subtitle":{"font":{"family":"Inter","size":"13pt","weight":800},"color":"#d8d8d8","line_height":23}},"vertical_padding":14},"body":{"fill_color":"'''
    obj_flow_template_header_74 = var_obj_flow_objects_property_box_colour
    obj_flow_template_header_75 = '''","stroke_color":"#383838","field_name":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#b3b3b3","padding":12},"field_value":{"font":{"family":"Inter","size":"10.5pt"},"color":"#bfbfbf","line_height":20,"padding":22},"vertical_padding":18},"select_outline":{"color":"#e6d845","padding":4,"border_radius":9},"anchor_markers":{"color":"'''
    obj_flow_template_header_76 = var_obj_flow_objects_property_anchor_markers_colour
    obj_flow_template_header_77 = '''","size":3},"border_radius":5,"horizontal_padding":20}},{"id":"observed_data","namespace":"stix_object.observed_data","type":2,"role":4096,"properties":{"first_observed":{"type":3,"is_required":true},"last_observed":{"type":3,"is_required":true},"number_observed":{"type":0,"min":0,"is_required":true}},"anchor_template":"@__builtin__anchor","style":{"max_width":320,"head":{"fill_color":"#737373","stroke_color":"#8c8c8c","one_title":{"title":{"font":{"family":"Inter","size":"10.5pt","weight":800},"color":"#d8d8d8"}},"two_title":{"title":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#d8d8d8","padding":8},"subtitle":{"font":{"family":"Inter","size":"13pt","weight":800},"color":"#d8d8d8","line_height":23}},"vertical_padding":14},"body":{"fill_color":"'''
    obj_flow_template_header_78 = var_obj_flow_objects_property_box_colour
    obj_flow_template_header_79 = '''","stroke_color":"#383838","field_name":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#b3b3b3","padding":12},"field_value":{"font":{"family":"Inter","size":"10.5pt"},"color":"#bfbfbf","line_height":20,"padding":22},"vertical_padding":18},"select_outline":{"color":"#e6d845","padding":4,"border_radius":9},"anchor_markers":{"color":"'''
    obj_flow_template_header_80 = var_obj_flow_objects_property_anchor_markers_colour
    obj_flow_template_header_81 = '''","size":3},"border_radius":5,"horizontal_padding":20}},{"id":"opinion","namespace":"stix_object.opinion","type":2,"role":4096,"properties":{"explanation":{"type":2,"is_primary":true},"authors":{"type":5,"form":{"type":2}},"opinion":{"type":2,"is_required":true}},"anchor_template":"@__builtin__anchor","style":{"max_width":320,"head":{"fill_color":"#737373","stroke_color":"#8c8c8c","one_title":{"title":{"font":{"family":"Inter","size":"10.5pt","weight":800},"color":"#d8d8d8"}},"two_title":{"title":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#d8d8d8","padding":8},"subtitle":{"font":{"family":"Inter","size":"13pt","weight":800},"color":"#d8d8d8","line_height":23}},"vertical_padding":14},"body":{"fill_color":"'''
    obj_flow_template_header_82 = var_obj_flow_objects_property_box_colour
    obj_flow_template_header_83 = '''","stroke_color":"#383838","field_name":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#b3b3b3","padding":12},"field_value":{"font":{"family":"Inter","size":"10.5pt"},"color":"#bfbfbf","line_height":20,"padding":22},"vertical_padding":18},"select_outline":{"color":"#e6d845","padding":4,"border_radius":9},"anchor_markers":{"color":"'''
    obj_flow_template_header_84 = var_obj_flow_objects_property_anchor_markers_colour
    obj_flow_template_header_85 = '''","size":3},"border_radius":5,"horizontal_padding":20}},{"id":"report","namespace":"stix_object.report","type":2,"role":4096,"properties":{"name":{"type":2,"is_primary":true,"is_required":true},"description":{"type":2},"report_types":{"type":5,"form":{"type":2,"is_required":true}},"published":{"type":3,"is_required":true}},"anchor_template":"@__builtin__anchor","style":{"max_width":320,"head":{"fill_color":"#737373","stroke_color":"#8c8c8c","one_title":{"title":{"font":{"family":"Inter","size":"10.5pt","weight":800},"color":"#d8d8d8"}},"two_title":{"title":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#d8d8d8","padding":8},"subtitle":{"font":{"family":"Inter","size":"13pt","weight":800},"color":"#d8d8d8","line_height":23}},"vertical_padding":14},"body":{"fill_color":"'''
    obj_flow_template_header_86 = var_obj_flow_objects_property_box_colour
    obj_flow_template_header_87 = '''","stroke_color":"#383838","field_name":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#b3b3b3","padding":12},"field_value":{"font":{"family":"Inter","size":"10.5pt"},"color":"#bfbfbf","line_height":20,"padding":22},"vertical_padding":18},"select_outline":{"color":"#e6d845","padding":4,"border_radius":9},"anchor_markers":{"color":"'''
    obj_flow_template_header_88 = var_obj_flow_objects_property_anchor_markers_colour
    obj_flow_template_header_89 = '''","size":3},"border_radius":5,"horizontal_padding":20}},{"id":"threat_actor","namespace":"stix_object.threat_actor","type":2,"role":4096,"properties":{"name":{"type":2,"is_primary":true,"is_required":true},"description":{"type":2},"threat_actor_types":{"type":5,"form":{"type":2,"is_required":true}},"aliases":{"type":5,"form":{"type":2}},"first_seen":{"type":3},"last_seen":{"type":3},"roles":{"type":5,"form":{"type":2}},"goals":{"type":5,"form":{"type":2}},"sophistication":{"type":2},"resource_level":{"type":2},"primary_motivation":{"type":2},"secondary_motivations":{"type":5,"form":{"type":2}},"personal_motivations":{"type":5,"form":{"type":2}}},"anchor_template":"@__builtin__anchor","style":{"max_width":320,"head":{"fill_color":"#737373","stroke_color":"#8c8c8c","one_title":{"title":{"font":{"family":"Inter","size":"10.5pt","weight":800},"color":"#d8d8d8"}},"two_title":{"title":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#d8d8d8","padding":8},"subtitle":{"font":{"family":"Inter","size":"13pt","weight":800},"color":"#d8d8d8","line_height":23}},"vertical_padding":14},"body":{"fill_color":"'''
    obj_flow_template_header_90 = var_obj_flow_objects_property_box_colour
    obj_flow_template_header_91 = '''","stroke_color":"#383838","field_name":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#b3b3b3","padding":12},"field_value":{"font":{"family":"Inter","size":"10.5pt"},"color":"#bfbfbf","line_height":20,"padding":22},"vertical_padding":18},"select_outline":{"color":"#e6d845","padding":4,"border_radius":9},"anchor_markers":{"color":"'''
    obj_flow_template_header_92 = var_obj_flow_objects_property_anchor_markers_colour
    obj_flow_template_header_93 = '''","size":3},"border_radius":5,"horizontal_padding":20}},{"id":"tool","namespace":"stix_object.tool","type":2,"role":4096,"properties":{"name":{"type":2,"is_primary":true,"is_required":true},"description":{"type":2},"tool_types":{"type":5,"form":{"type":2,"is_required":true}},"aliases":{"type":5,"form":{"type":2}},"kill_chain_phases":{"type":5,"form":{"type":2}},"tool_version":{"type":2}},"anchor_template":"@__builtin__anchor","style":{"max_width":320,"head":{"fill_color":"#737373","stroke_color":"#8c8c8c","one_title":{"title":{"font":{"family":"Inter","size":"10.5pt","weight":800},"color":"#d8d8d8"}},"two_title":{"title":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#d8d8d8","padding":8},"subtitle":{"font":{"family":"Inter","size":"13pt","weight":800},"color":"#d8d8d8","line_height":23}},"vertical_padding":14},"body":{"fill_color":"'''
    obj_flow_template_header_94 = var_obj_flow_objects_property_box_colour
    obj_flow_template_header_95 = '''","stroke_color":"#383838","field_name":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#b3b3b3","padding":12},"field_value":{"font":{"family":"Inter","size":"10.5pt"},"color":"#bfbfbf","line_height":20,"padding":22},"vertical_padding":18},"select_outline":{"color":"#e6d845","padding":4,"border_radius":9},"anchor_markers":{"color":"'''
    obj_flow_template_header_96 = var_obj_flow_objects_property_anchor_markers_colour
    obj_flow_template_header_97 = '''","size":3},"border_radius":5,"horizontal_padding":20}},{"id":"vulnerability","namespace":"stix_object.vulnerability","type":2,"role":4096,"properties":{"name":{"type":2,"is_primary":true,"is_required":true},"description":{"type":2}},"anchor_template":"@__builtin__anchor","style":{"max_width":320,"head":{"fill_color":"#737373","stroke_color":"#8c8c8c","one_title":{"title":{"font":{"family":"Inter","size":"10.5pt","weight":800},"color":"#d8d8d8"}},"two_title":{"title":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#d8d8d8","padding":8},"subtitle":{"font":{"family":"Inter","size":"13pt","weight":800},"color":"#d8d8d8","line_height":23}},"vertical_padding":14},"body":{"fill_color":"'''
    obj_flow_template_header_98 = var_obj_flow_objects_property_box_colour
    obj_flow_template_header_99 = '''","stroke_color":"#383838","field_name":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#b3b3b3","padding":12},"field_value":{"font":{"family":"Inter","size":"10.5pt"},"color":"#bfbfbf","line_height":20,"padding":22},"vertical_padding":18},"select_outline":{"color":"#e6d845","padding":4,"border_radius":9},"anchor_markers":{"color":"'''
    obj_flow_template_header_100 = var_obj_flow_objects_property_anchor_markers_colour
    obj_flow_template_header_101 = '''","size":3},"border_radius":5,"horizontal_padding":20}},{"id":"artifact","namespace":"stix_observable.artifact","type":2,"role":4096,"properties":{"mime_type":{"type":2},"payload_bin":{"type":2},"url":{"type":2},"hashes":{"type":2},"encryption_algorithm":{"type":2},"decryption_key":{"type":2}},"anchor_template":"@__builtin__anchor","style":{"max_width":320,"head":{"fill_color":"#737373","stroke_color":"#8c8c8c","one_title":{"title":{"font":{"family":"Inter","size":"10.5pt","weight":800},"color":"#d8d8d8"}},"two_title":{"title":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#d8d8d8","padding":8},"subtitle":{"font":{"family":"Inter","size":"13pt","weight":800},"color":"#d8d8d8","line_height":23}},"vertical_padding":14},"body":{"fill_color":"'''
    obj_flow_template_header_102 = var_obj_flow_objects_property_box_colour
    obj_flow_template_header_103 = '''","stroke_color":"#383838","field_name":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#b3b3b3","padding":12},"field_value":{"font":{"family":"Inter","size":"10.5pt"},"color":"#bfbfbf","line_height":20,"padding":22},"vertical_padding":18},"select_outline":{"color":"#e6d845","padding":4,"border_radius":9},"anchor_markers":{"color":"'''
    obj_flow_template_header_104 = var_obj_flow_objects_property_anchor_markers_colour
    obj_flow_template_header_105 = '''","size":3},"border_radius":5,"horizontal_padding":20}},{"id":"autonomous_system","namespace":"stix_observable.autonomous_system","type":2,"role":4096,"properties":{"number":{"type":2,"is_primary":true,"is_required":true},"name":{"type":2},"rir":{"type":2}},"anchor_template":"@__builtin__anchor","style":{"max_width":320,"head":{"fill_color":"#737373","stroke_color":"#8c8c8c","one_title":{"title":{"font":{"family":"Inter","size":"10.5pt","weight":800},"color":"#d8d8d8"}},"two_title":{"title":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#d8d8d8","padding":8},"subtitle":{"font":{"family":"Inter","size":"13pt","weight":800},"color":"#d8d8d8","line_height":23}},"vertical_padding":14},"body":{"fill_color":"'''
    obj_flow_template_header_106 = var_obj_flow_objects_property_box_colour
    obj_flow_template_header_107 = '''","stroke_color":"#383838","field_name":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#b3b3b3","padding":12},"field_value":{"font":{"family":"Inter","size":"10.5pt"},"color":"#bfbfbf","line_height":20,"padding":22},"vertical_padding":18},"select_outline":{"color":"#e6d845","padding":4,"border_radius":9},"anchor_markers":{"color":"'''
    obj_flow_template_header_108 = var_obj_flow_objects_property_anchor_markers_colour
    obj_flow_template_header_109 = '''","size":3},"border_radius":5,"horizontal_padding":20}},{"id":"directory","namespace":"stix_observable.directory","type":2,"role":4096,"properties":{"path":{"type":2,"is_primary":true,"is_required":true},"path_enc":{"type":2},"ctime":{"type":3},"mtime":{"type":3},"atime":{"type":3}},"anchor_template":"@__builtin__anchor","style":{"max_width":320,"head":{"fill_color":"#737373","stroke_color":"#8c8c8c","one_title":{"title":{"font":{"family":"Inter","size":"10.5pt","weight":800},"color":"#d8d8d8"}},"two_title":{"title":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#d8d8d8","padding":8},"subtitle":{"font":{"family":"Inter","size":"13pt","weight":800},"color":"#d8d8d8","line_height":23}},"vertical_padding":14},"body":{"fill_color":"'''
    obj_flow_template_header_110 = var_obj_flow_objects_property_box_colour
    obj_flow_template_header_111 = '''","stroke_color":"#383838","field_name":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#b3b3b3","padding":12},"field_value":{"font":{"family":"Inter","size":"10.5pt"},"color":"#bfbfbf","line_height":20,"padding":22},"vertical_padding":18},"select_outline":{"color":"#e6d845","padding":4,"border_radius":9},"anchor_markers":{"color":"'''
    obj_flow_template_header_112 = var_obj_flow_objects_property_anchor_markers_colour
    obj_flow_template_header_113 = '''","size":3},"border_radius":5,"horizontal_padding":20}},{"id":"domain_name","namespace":"stix_observable.domain_name","type":2,"role":4096,"properties":{"value":{"type":2,"is_required":true}},"anchor_template":"@__builtin__anchor","style":{"max_width":320,"head":{"fill_color":"#737373","stroke_color":"#8c8c8c","one_title":{"title":{"font":{"family":"Inter","size":"10.5pt","weight":800},"color":"#d8d8d8"}},"two_title":{"title":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#d8d8d8","padding":8},"subtitle":{"font":{"family":"Inter","size":"13pt","weight":800},"color":"#d8d8d8","line_height":23}},"vertical_padding":14},"body":{"fill_color":"'''
    obj_flow_template_header_114 = var_obj_flow_objects_property_box_colour
    obj_flow_template_header_115 = '''","stroke_color":"#383838","field_name":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#b3b3b3","padding":12},"field_value":{"font":{"family":"Inter","size":"10.5pt"},"color":"#bfbfbf","line_height":20,"padding":22},"vertical_padding":18},"select_outline":{"color":"#e6d845","padding":4,"border_radius":9},"anchor_markers":{"color":"'''
    obj_flow_template_header_116 = var_obj_flow_objects_property_anchor_markers_colour
    obj_flow_template_header_117 = '''","size":3},"border_radius":5,"horizontal_padding":20}},{"id":"email_address","namespace":"stix_observable.email_address","type":2,"role":4096,"properties":{"value":{"type":2,"is_required":true},"display_name":{"type":2}},"anchor_template":"@__builtin__anchor","style":{"max_width":320,"head":{"fill_color":"#737373","stroke_color":"#8c8c8c","one_title":{"title":{"font":{"family":"Inter","size":"10.5pt","weight":800},"color":"#d8d8d8"}},"two_title":{"title":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#d8d8d8","padding":8},"subtitle":{"font":{"family":"Inter","size":"13pt","weight":800},"color":"#d8d8d8","line_height":23}},"vertical_padding":14},"body":{"fill_color":"'''
    obj_flow_template_header_118 = var_obj_flow_objects_property_box_colour
    obj_flow_template_header_119 = '''","stroke_color":"#383838","field_name":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#b3b3b3","padding":12},"field_value":{"font":{"family":"Inter","size":"10.5pt"},"color":"#bfbfbf","line_height":20,"padding":22},"vertical_padding":18},"select_outline":{"color":"#e6d845","padding":4,"border_radius":9},"anchor_markers":{"color":"'''
    obj_flow_template_header_120 = var_obj_flow_objects_property_anchor_markers_colour
    obj_flow_template_header_121 = '''","size":3},"border_radius":5,"horizontal_padding":20}},{"id":"email_message","namespace":"stix_observable.email_message","type":2,"role":4096,"properties":{"is_multipart":{"type":4,"options":{"type":5,"form":{"type":2},"value":[["true","True"],["false","False"]]},"is_required":true},"date":{"type":2},"content_type":{"type":2},"message_id":{"type":2},"subject":{"type":2,"is_primary":true},"received_lines":{"type":2},"additional_header_fields":{"type":2},"body":{"type":2},"body_multipart":{"type":2}},"anchor_template":"@__builtin__anchor","style":{"max_width":320,"head":{"fill_color":"#737373","stroke_color":"#8c8c8c","one_title":{"title":{"font":{"family":"Inter","size":"10.5pt","weight":800},"color":"#d8d8d8"}},"two_title":{"title":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#d8d8d8","padding":8},"subtitle":{"font":{"family":"Inter","size":"13pt","weight":800},"color":"#d8d8d8","line_height":23}},"vertical_padding":14},"body":{"fill_color":"'''
    obj_flow_template_header_122 = var_obj_flow_objects_property_box_colour
    obj_flow_template_header_123 = '''","stroke_color":"#383838","field_name":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#b3b3b3","padding":12},"field_value":{"font":{"family":"Inter","size":"10.5pt"},"color":"#bfbfbf","line_height":20,"padding":22},"vertical_padding":18},"select_outline":{"color":"#e6d845","padding":4,"border_radius":9},"anchor_markers":{"color":"'''
    obj_flow_template_header_124 = var_obj_flow_objects_property_anchor_markers_colour
    obj_flow_template_header_125 = '''","size":3},"border_radius":5,"horizontal_padding":20}},{"id":"file","namespace":"stix_observable.file","type":2,"role":4096,"properties":{"hashes":{"type":2},"size":{"type":2},"name":{"type":2,"is_primary":true},"name_enc":{"type":2},"magic_number_hex":{"type":2},"mime_type":{"type":2},"ctime":{"type":3},"mtime":{"type":3},"atime":{"type":3}},"anchor_template":"@__builtin__anchor","style":{"max_width":320,"head":{"fill_color":"#737373","stroke_color":"#8c8c8c","one_title":{"title":{"font":{"family":"Inter","size":"10.5pt","weight":800},"color":"#d8d8d8"}},"two_title":{"title":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#d8d8d8","padding":8},"subtitle":{"font":{"family":"Inter","size":"13pt","weight":800},"color":"#d8d8d8","line_height":23}},"vertical_padding":14},"body":{"fill_color":"'''
    obj_flow_template_header_126 = var_obj_flow_objects_property_box_colour
    obj_flow_template_header_127 = '''","stroke_color":"#383838","field_name":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#b3b3b3","padding":12},"field_value":{"font":{"family":"Inter","size":"10.5pt"},"color":"#bfbfbf","line_height":20,"padding":22},"vertical_padding":18},"select_outline":{"color":"#e6d845","padding":4,"border_radius":9},"anchor_markers":{"color":"'''
    obj_flow_template_header_128 = var_obj_flow_objects_property_anchor_markers_colour
    obj_flow_template_header_129 = '''","size":3},"border_radius":5,"horizontal_padding":20}},{"id":"ipv4_addr","namespace":"stix_observable.ipv4_addr","type":2,"role":4096,"properties":{"value":{"type":2,"is_required":true}},"anchor_template":"@__builtin__anchor","style":{"max_width":320,"head":{"fill_color":"#737373","stroke_color":"#8c8c8c","one_title":{"title":{"font":{"family":"Inter","size":"10.5pt","weight":800},"color":"#d8d8d8"}},"two_title":{"title":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#d8d8d8","padding":8},"subtitle":{"font":{"family":"Inter","size":"13pt","weight":800},"color":"#d8d8d8","line_height":23}},"vertical_padding":14},"body":{"fill_color":"'''
    obj_flow_template_header_130 = var_obj_flow_objects_property_box_colour
    obj_flow_template_header_131 = '''","stroke_color":"#383838","field_name":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#b3b3b3","padding":12},"field_value":{"font":{"family":"Inter","size":"10.5pt"},"color":"#bfbfbf","line_height":20,"padding":22},"vertical_padding":18},"select_outline":{"color":"#e6d845","padding":4,"border_radius":9},"anchor_markers":{"color":"'''
    obj_flow_template_header_132 = var_obj_flow_objects_property_anchor_markers_colour
    obj_flow_template_header_133 = '''","size":3},"border_radius":5,"horizontal_padding":20}},{"id":"ipv6_addr","namespace":"stix_observable.ipv6_addr","type":2,"role":4096,"properties":{"value":{"type":2,"is_required":true}},"anchor_template":"@__builtin__anchor","style":{"max_width":320,"head":{"fill_color":"#737373","stroke_color":"#8c8c8c","one_title":{"title":{"font":{"family":"Inter","size":"10.5pt","weight":800},"color":"#d8d8d8"}},"two_title":{"title":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#d8d8d8","padding":8},"subtitle":{"font":{"family":"Inter","size":"13pt","weight":800},"color":"#d8d8d8","line_height":23}},"vertical_padding":14},"body":{"fill_color":"'''
    obj_flow_template_header_134 = var_obj_flow_objects_property_box_colour
    obj_flow_template_header_135 = '''","stroke_color":"#383838","field_name":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#b3b3b3","padding":12},"field_value":{"font":{"family":"Inter","size":"10.5pt"},"color":"#bfbfbf","line_height":20,"padding":22},"vertical_padding":18},"select_outline":{"color":"#e6d845","padding":4,"border_radius":9},"anchor_markers":{"color":"'''
    obj_flow_template_header_136 = var_obj_flow_objects_property_anchor_markers_colour
    obj_flow_template_header_137 = '''","size":3},"border_radius":5,"horizontal_padding":20}},{"id":"mac_addr","namespace":"stix_observable.mac_addr","type":2,"role":4096,"properties":{"value":{"type":2,"is_required":true}},"anchor_template":"@__builtin__anchor","style":{"max_width":320,"head":{"fill_color":"#737373","stroke_color":"#8c8c8c","one_title":{"title":{"font":{"family":"Inter","size":"10.5pt","weight":800},"color":"#d8d8d8"}},"two_title":{"title":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#d8d8d8","padding":8},"subtitle":{"font":{"family":"Inter","size":"13pt","weight":800},"color":"#d8d8d8","line_height":23}},"vertical_padding":14},"body":{"fill_color":"'''
    obj_flow_template_header_138 = var_obj_flow_objects_property_box_colour
    obj_flow_template_header_139 = '''","stroke_color":"#383838","field_name":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#b3b3b3","padding":12},"field_value":{"font":{"family":"Inter","size":"10.5pt"},"color":"#bfbfbf","line_height":20,"padding":22},"vertical_padding":18},"select_outline":{"color":"#e6d845","padding":4,"border_radius":9},"anchor_markers":{"color":"'''
    obj_flow_template_header_140 = var_obj_flow_objects_property_anchor_markers_colour
    obj_flow_template_header_141 = '''","size":3},"border_radius":5,"horizontal_padding":20}},{"id":"mutex","namespace":"stix_observable.mutex","type":2,"role":4096,"properties":{"name":{"type":2,"is_required":true}},"anchor_template":"@__builtin__anchor","style":{"max_width":320,"head":{"fill_color":"#737373","stroke_color":"#8c8c8c","one_title":{"title":{"font":{"family":"Inter","size":"10.5pt","weight":800},"color":"#d8d8d8"}},"two_title":{"title":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#d8d8d8","padding":8},"subtitle":{"font":{"family":"Inter","size":"13pt","weight":800},"color":"#d8d8d8","line_height":23}},"vertical_padding":14},"body":{"fill_color":"'''
    obj_flow_template_header_142 = var_obj_flow_objects_property_box_colour
    obj_flow_template_header_143 = '''","stroke_color":"#383838","field_name":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#b3b3b3","padding":12},"field_value":{"font":{"family":"Inter","size":"10.5pt"},"color":"#bfbfbf","line_height":20,"padding":22},"vertical_padding":18},"select_outline":{"color":"#e6d845","padding":4,"border_radius":9},"anchor_markers":{"color":"'''
    obj_flow_template_header_144 = var_obj_flow_objects_property_anchor_markers_colour
    obj_flow_template_header_145 = '''","size":3},"border_radius":5,"horizontal_padding":20}},{"id":"network_traffic","namespace":"stix_observable.network_traffic","type":2,"role":4096,"properties":{"start":{"type":3},"end":{"type":3},"is_active":{"type":4,"options":{"type":5,"form":{"type":2},"value":[["true","True"],["false","False"]]}},"src_port":{"type":0,"min":0,"max":65535},"dst_port":{"type":0,"min":0,"max":65535},"protocols":{"type":5,"form":{"type":2,"is_required":true}},"src_byte_count":{"type":0,"min":0},"dst_byte_count":{"type":0,"min":0},"src_packets":{"type":0,"min":0},"dst_packets":{"type":0,"min":0},"ipfix":{"type":2}},"anchor_template":"@__builtin__anchor","style":{"max_width":320,"head":{"fill_color":"#737373","stroke_color":"#8c8c8c","one_title":{"title":{"font":{"family":"Inter","size":"10.5pt","weight":800},"color":"#d8d8d8"}},"two_title":{"title":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#d8d8d8","padding":8},"subtitle":{"font":{"family":"Inter","size":"13pt","weight":800},"color":"#d8d8d8","line_height":23}},"vertical_padding":14},"body":{"fill_color":"'''
    obj_flow_template_header_146 = var_obj_flow_objects_property_box_colour
    obj_flow_template_header_147 = '''","stroke_color":"#383838","field_name":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#b3b3b3","padding":12},"field_value":{"font":{"family":"Inter","size":"10.5pt"},"color":"#bfbfbf","line_height":20,"padding":22},"vertical_padding":18},"select_outline":{"color":"#e6d845","padding":4,"border_radius":9},"anchor_markers":{"color":"'''
    obj_flow_template_header_148 = var_obj_flow_objects_property_anchor_markers_colour
    obj_flow_template_header_149 = '''","size":3},"border_radius":5,"horizontal_padding":20}},{"id":"process","namespace":"stix_observable.process","type":2,"role":4096,"properties":{"is_hidden":{"type":4,"options":{"type":5,"form":{"type":2},"value":[["true","True"],["false","False"]]}},"pid":{"type":0,"min":0},"created_time":{"type":3},"cwd":{"type":2},"command_line":{"type":2,"is_required":true},"environment_variables":{"type":2}},"anchor_template":"@__builtin__anchor","style":{"max_width":320,"head":{"fill_color":"#737373","stroke_color":"#8c8c8c","one_title":{"title":{"font":{"family":"Inter","size":"10.5pt","weight":800},"color":"#d8d8d8"}},"two_title":{"title":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#d8d8d8","padding":8},"subtitle":{"font":{"family":"Inter","size":"13pt","weight":800},"color":"#d8d8d8","line_height":23}},"vertical_padding":14},"body":{"fill_color":"'''
    obj_flow_template_header_150 = var_obj_flow_objects_property_box_colour
    obj_flow_template_header_151 = '''","stroke_color":"#383838","field_name":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#b3b3b3","padding":12},"field_value":{"font":{"family":"Inter","size":"10.5pt"},"color":"#bfbfbf","line_height":20,"padding":22},"vertical_padding":18},"select_outline":{"color":"#e6d845","padding":4,"border_radius":9},"anchor_markers":{"color":"'''
    obj_flow_template_header_152 = var_obj_flow_objects_property_anchor_markers_colour
    obj_flow_template_header_153 = '''","size":3},"border_radius":5,"horizontal_padding":20}},{"id":"software","namespace":"stix_observable.software","type":2,"role":4096,"properties":{"name":{"type":2,"is_primary":true,"is_required":true},"cpe":{"type":2},"languages":{"type":5,"form":{"type":2}},"vendor":{"type":2},"version":{"type":2}},"anchor_template":"@__builtin__anchor","style":{"max_width":320,"head":{"fill_color":"#737373","stroke_color":"#8c8c8c","one_title":{"title":{"font":{"family":"Inter","size":"10.5pt","weight":800},"color":"#d8d8d8"}},"two_title":{"title":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#d8d8d8","padding":8},"subtitle":{"font":{"family":"Inter","size":"13pt","weight":800},"color":"#d8d8d8","line_height":23}},"vertical_padding":14},"body":{"fill_color":"'''
    obj_flow_template_header_154 = var_obj_flow_objects_property_box_colour
    obj_flow_template_header_155 = '''","stroke_color":"#383838","field_name":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#b3b3b3","padding":12},"field_value":{"font":{"family":"Inter","size":"10.5pt"},"color":"#bfbfbf","line_height":20,"padding":22},"vertical_padding":18},"select_outline":{"color":"#e6d845","padding":4,"border_radius":9},"anchor_markers":{"color":"'''
    obj_flow_template_header_156 = var_obj_flow_objects_property_anchor_markers_colour
    obj_flow_template_header_157 = '''","size":3},"border_radius":5,"horizontal_padding":20}},{"id":"url","namespace":"stix_observable.url","type":2,"role":4096,"properties":{"value":{"type":2,"is_required":true}},"anchor_template":"@__builtin__anchor","style":{"max_width":320,"head":{"fill_color":"#737373","stroke_color":"#8c8c8c","one_title":{"title":{"font":{"family":"Inter","size":"10.5pt","weight":800},"color":"#d8d8d8"}},"two_title":{"title":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#d8d8d8","padding":8},"subtitle":{"font":{"family":"Inter","size":"13pt","weight":800},"color":"#d8d8d8","line_height":23}},"vertical_padding":14},"body":{"fill_color":"'''
    obj_flow_template_header_158 = var_obj_flow_objects_property_box_colour
    obj_flow_template_header_159 = '''","stroke_color":"#383838","field_name":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#b3b3b3","padding":12},"field_value":{"font":{"family":"Inter","size":"10.5pt"},"color":"#bfbfbf","line_height":20,"padding":22},"vertical_padding":18},"select_outline":{"color":"#e6d845","padding":4,"border_radius":9},"anchor_markers":{"color":"'''
    obj_flow_template_header_160 = var_obj_flow_objects_property_anchor_markers_colour
    obj_flow_template_header_161 = '''","size":3},"border_radius":5,"horizontal_padding":20}},{"id":"user_account","namespace":"stix_observable.user_account","type":2,"role":4096,"properties":{"user_id":{"type":2},"credential":{"type":2},"account_login":{"type":2},"account_type":{"type":2},"display_name":{"type":2,"is_primary":true,"is_required":true},"is_service_account":{"type":4,"options":{"type":5,"form":{"type":2},"value":[["true","True"],["false","False"]]}},"is_privileged":{"type":4,"options":{"type":5,"form":{"type":2},"value":[["true","True"],["false","False"]]}},"can_escalate_privs":{"type":4,"options":{"type":5,"form":{"type":2},"value":[["true","True"],["false","False"]]}},"is_disabled":{"type":4,"options":{"type":5,"form":{"type":2},"value":[["true","True"],["false","False"]]}},"account_created":{"type":3},"account_expires":{"type":3},"credential_last_changed":{"type":3},"account_first_login":{"type":3},"account_last_login":{"type":3}},"anchor_template":"@__builtin__anchor","style":{"max_width":320,"head":{"fill_color":"#737373","stroke_color":"#8c8c8c","one_title":{"title":{"font":{"family":"Inter","size":"10.5pt","weight":800},"color":"#d8d8d8"}},"two_title":{"title":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#d8d8d8","padding":8},"subtitle":{"font":{"family":"Inter","size":"13pt","weight":800},"color":"#d8d8d8","line_height":23}},"vertical_padding":14},"body":{"fill_color":"'''
    obj_flow_template_header_162 = var_obj_flow_objects_property_box_colour
    obj_flow_template_header_163 = '''","stroke_color":"#383838","field_name":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#b3b3b3","padding":12},"field_value":{"font":{"family":"Inter","size":"10.5pt"},"color":"#bfbfbf","line_height":20,"padding":22},"vertical_padding":18},"select_outline":{"color":"#e6d845","padding":4,"border_radius":9},"anchor_markers":{"color":"'''
    obj_flow_template_header_164 = var_obj_flow_objects_property_anchor_markers_colour
    obj_flow_template_header_165 = '''","size":3},"border_radius":5,"horizontal_padding":20}},{"id":"windows_registry_key","namespace":"stix_observable.windows_registry_key","type":2,"role":4096,"properties":{"key":{"type":2,"is_primary":true},"values":{"type":5,"form":{"type":2}},"modified_time":{"type":3},"number_of_subkeys":{"type":0,"min":0}},"anchor_template":"@__builtin__anchor","style":{"max_width":320,"head":{"fill_color":"#737373","stroke_color":"#8c8c8c","one_title":{"title":{"font":{"family":"Inter","size":"10.5pt","weight":800},"color":"#d8d8d8"}},"two_title":{"title":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#d8d8d8","padding":8},"subtitle":{"font":{"family":"Inter","size":"13pt","weight":800},"color":"#d8d8d8","line_height":23}},"vertical_padding":14},"body":{"fill_color":"'''
    obj_flow_template_header_166 = var_obj_flow_objects_property_box_colour
    obj_flow_template_header_167 = '''","stroke_color":"#383838","field_name":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#b3b3b3","padding":12},"field_value":{"font":{"family":"Inter","size":"10.5pt"},"color":"#bfbfbf","line_height":20,"padding":22},"vertical_padding":18},"select_outline":{"color":"#e6d845","padding":4,"border_radius":9},"anchor_markers":{"color":"'''
    obj_flow_template_header_168 = var_obj_flow_objects_property_anchor_markers_colour
    obj_flow_template_header_169 = '''","size":3},"border_radius":5,"horizontal_padding":20}},{"id":"x509_certificate","namespace":"stix_observable.x509_certificate","type":2,"role":4096,"properties":{"subject":{"type":2,"is_primary":true,"is_required":true},"is_self_signed":{"type":4,"options":{"type":5,"form":{"type":2},"value":[["true","True"],["false","False"]]}},"hashes":{"type":2},"version":{"type":2},"serial_number":{"type":2},"signature_algorithm":{"type":2},"issuer":{"type":2},"validity_not_before":{"type":3},"validity_not_after":{"type":3},"subject_public_key_algorithm":{"type":2},"subject_public_key_modulus":{"type":2},"subject_public_key_exponent":{"type":0,"min":0}},"anchor_template":"@__builtin__anchor","style":{"max_width":320,"head":{"fill_color":"#737373","stroke_color":"#8c8c8c","one_title":{"title":{"font":{"family":"Inter","size":"10.5pt","weight":800},"color":"#d8d8d8"}},"two_title":{"title":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#d8d8d8","padding":8},"subtitle":{"font":{"family":"Inter","size":"13pt","weight":800},"color":"#d8d8d8","line_height":23}},"vertical_padding":14},"body":{"fill_color":"'''
    obj_flow_template_header_170 = var_obj_flow_objects_property_box_colour
    obj_flow_template_header_171 = '''","stroke_color":"#383838","field_name":{"font":{"family":"Inter","size":"8pt","weight":600},"color":"#b3b3b3","padding":12},"field_value":{"font":{"family":"Inter","size":"10.5pt"},"color":"#bfbfbf","line_height":20,"padding":22},"vertical_padding":18},"select_outline":{"color":"#e6d845","padding":4,"border_radius":9},"anchor_markers":{"color":"'''
    obj_flow_template_header_172 = var_obj_flow_objects_property_anchor_markers_colour
    obj_flow_template_header_173 = '''","size":3},"border_radius":5,"horizontal_padding":20}}]},'''
    obj_flow_template_header = obj_flow_template_header_0 + obj_flow_template_header_1 + obj_flow_template_header_2 + obj_flow_template_header_3 + obj_flow_template_header_4 + obj_flow_template_header_5 + obj_flow_template_header_6 + obj_flow_template_header_7 + obj_flow_template_header_8 + obj_flow_template_header_9 + obj_flow_template_header_10 + obj_flow_template_header_11 + obj_flow_template_header_12 + obj_flow_template_header_13 + obj_flow_template_header_14 + obj_flow_template_header_15 + obj_flow_template_header_16 + obj_flow_template_header_17 + obj_flow_template_header_18 + obj_flow_template_header_19 + obj_flow_template_header_20 + obj_flow_template_header_21 + obj_flow_template_header_22 + obj_flow_template_header_23 + obj_flow_template_header_24 + obj_flow_template_header_25 + obj_flow_template_header_26 + obj_flow_template_header_27 + obj_flow_template_header_28 + obj_flow_template_header_29 + obj_flow_template_header_30 + obj_flow_template_header_31 + obj_flow_template_header_32 + obj_flow_template_header_33 + obj_flow_template_header_34 + obj_flow_template_header_35 + obj_flow_template_header_36 + obj_flow_template_header_37 + obj_flow_template_header_38 + obj_flow_template_header_39 + obj_flow_template_header_40 + obj_flow_template_header_41 + obj_flow_template_header_42 + obj_flow_template_header_43 + obj_flow_template_header_44 + obj_flow_template_header_45 + obj_flow_template_header_46 + obj_flow_template_header_47 + obj_flow_template_header_48 + obj_flow_template_header_49 + obj_flow_template_header_50 + obj_flow_template_header_51 + obj_flow_template_header_52 + obj_flow_template_header_53 + obj_flow_template_header_54 + obj_flow_template_header_55 + obj_flow_template_header_56 + obj_flow_template_header_57 + obj_flow_template_header_58 + obj_flow_template_header_59 + obj_flow_template_header_60 + obj_flow_template_header_61 + obj_flow_template_header_62 + obj_flow_template_header_63 + obj_flow_template_header_64 + obj_flow_template_header_65 + obj_flow_template_header_66 + obj_flow_template_header_67 + obj_flow_template_header_68 + obj_flow_template_header_69 + obj_flow_template_header_70 + obj_flow_template_header_71 + obj_flow_template_header_72 + obj_flow_template_header_73 + obj_flow_template_header_74 + obj_flow_template_header_75 + obj_flow_template_header_76 + obj_flow_template_header_77 + obj_flow_template_header_78 + obj_flow_template_header_79 + obj_flow_template_header_80 + obj_flow_template_header_81 + obj_flow_template_header_82 + obj_flow_template_header_83 + obj_flow_template_header_84 + obj_flow_template_header_85 + obj_flow_template_header_86 + obj_flow_template_header_87 + obj_flow_template_header_88 + obj_flow_template_header_89 + obj_flow_template_header_90 + obj_flow_template_header_91 + obj_flow_template_header_92 + obj_flow_template_header_93 + obj_flow_template_header_94 + obj_flow_template_header_95 + obj_flow_template_header_96 + obj_flow_template_header_97 + obj_flow_template_header_98 + obj_flow_template_header_99 + obj_flow_template_header_100 + obj_flow_template_header_101 + obj_flow_template_header_102 + obj_flow_template_header_103 + obj_flow_template_header_104 + obj_flow_template_header_105 + obj_flow_template_header_106 + obj_flow_template_header_107 + obj_flow_template_header_108 + obj_flow_template_header_109 + obj_flow_template_header_110 + obj_flow_template_header_111 + obj_flow_template_header_112 + obj_flow_template_header_113 + obj_flow_template_header_114 + obj_flow_template_header_115 + obj_flow_template_header_116 + obj_flow_template_header_117 + obj_flow_template_header_118 + obj_flow_template_header_119 + obj_flow_template_header_120 + obj_flow_template_header_121 + obj_flow_template_header_122 + obj_flow_template_header_123 + obj_flow_template_header_124 + obj_flow_template_header_125 + obj_flow_template_header_126 + obj_flow_template_header_127 + obj_flow_template_header_128 + obj_flow_template_header_129 + obj_flow_template_header_130 + obj_flow_template_header_131 + obj_flow_template_header_132 + obj_flow_template_header_133 + obj_flow_template_header_134 + obj_flow_template_header_135 + obj_flow_template_header_136 + obj_flow_template_header_137 + obj_flow_template_header_138 + obj_flow_template_header_139 + obj_flow_template_header_140 + obj_flow_template_header_141 + obj_flow_template_header_142 + obj_flow_template_header_143 + obj_flow_template_header_144 + obj_flow_template_header_145 + obj_flow_template_header_146 + obj_flow_template_header_147 + obj_flow_template_header_148 + obj_flow_template_header_149 + obj_flow_template_header_150 + obj_flow_template_header_151 + obj_flow_template_header_152 + obj_flow_template_header_153 + obj_flow_template_header_154 + obj_flow_template_header_155 + obj_flow_template_header_156 + obj_flow_template_header_157 + obj_flow_template_header_158 + obj_flow_template_header_159 + obj_flow_template_header_160 + obj_flow_template_header_161 + obj_flow_template_header_162 + obj_flow_template_header_163 + obj_flow_template_header_164 + obj_flow_template_header_165 + obj_flow_template_header_166 + obj_flow_template_header_167 + obj_flow_template_header_168 + obj_flow_template_header_169 + obj_flow_template_header_170 + obj_flow_template_header_171 + obj_flow_template_header_172 + obj_flow_template_header_173
    obj_flow_objects_actions_content = ""
    obj_flow_action_child_definition = ""
    obj_flow_action_child_header = ""
    obj_flow_objects_actions_child_content = ""
    flow_name_content = document_prefix_content
    var_x_pos = -290
    var_y_pos = -170
    for technique in array_obj_sorted_construct:
        obj_flow_action_child_header_GUID = technique['guid']
        obj_list_flow_objects_action_child_GUID = []
        for count in range(1, 13):
            obj_action_child_GUID = str(uuid.uuid4())
            obj_list_flow_objects_action_child_GUID.append(obj_action_child_GUID)
        attack_flow_action_child_GUID_group = '","'.join(str(guid) for guid in obj_list_flow_objects_action_child_GUID)
        obj_flow_action_child_header = f'''{{"id":"{obj_flow_action_child_header_GUID}","x":{var_x_pos},"y":{var_y_pos},"attrs":256,"template":"action","children":["{attack_flow_action_child_GUID_group}"],"properties":[["name","{technique["attack_name"]}"],["tactic_id",null],["tactic_ref","{technique["attack_tactics"][0]}"],["technique_id","{technique["attack_id"]}"],["technique_ref",null],["description","DESCRIPTION_PLACEHOLDER"],["confidence","62814720b26c68ab20bbb6669a1ec919"],["execution_start",null],["execution_end",null]]}},'''
        obj_flow_action_child_definition = f'''{{"id":"{obj_list_flow_objects_action_child_GUID[0]}","x":0,"y":0,"attrs":0,"template":"@__builtin__anchor","children":[],"properties":[],"angle":1}},{{"id":"{obj_list_flow_objects_action_child_GUID[1]}","x":0,"y":0,"attrs":0,"template":"@__builtin__anchor","children":[],"properties":[],"angle":1}},{{"id":"{obj_list_flow_objects_action_child_GUID[2]}","x":0,"y":0,"attrs":0,"template":"@__builtin__anchor","children":[],"properties":[],"angle":1}},{{"id":"{obj_list_flow_objects_action_child_GUID[3]}","x":0,"y":0,"attrs":0,"template":"@__builtin__anchor","children":[],"properties":[],"angle":0}},{{"id":"{obj_list_flow_objects_action_child_GUID[4]}","x":0,"y":0,"attrs":0,"template":"@__builtin__anchor","children":[],"properties":[],"angle":0}},{{"id":"{obj_list_flow_objects_action_child_GUID[5]}","x":0,"y":0,"attrs":0,"template":"@__builtin__anchor","children":[],"properties":[],"angle":0}},{{"id":"{obj_list_flow_objects_action_child_GUID[6]}","x":0,"y":0,"attrs":0,"template":"@__builtin__anchor","children":[],"properties":[],"angle":1}},{{"id":"{obj_list_flow_objects_action_child_GUID[7]}","x":0,"y":0,"attrs":0,"template":"@__builtin__anchor","children":[],"properties":[],"angle":1}},{{"id":"{obj_list_flow_objects_action_child_GUID[8]}","x":0,"y":0,"attrs":0,"template":"@__builtin__anchor","children":[],"properties":[],"angle":1}},{{"id":"{obj_list_flow_objects_action_child_GUID[9]}","x":0,"y":0,"attrs":0,"template":"@__builtin__anchor","children":[],"properties":[],"angle":0}},{{"id":"{obj_list_flow_objects_action_child_GUID[10]}","x":0,"y":0,"attrs":0,"template":"@__builtin__anchor","children":[],"properties":[],"angle":0}},{{"id":"{obj_list_flow_objects_action_child_GUID[11]}","x":0,"y":0,"attrs":0,"template":"@__builtin__anchor","children":[],"properties":[],"angle":0}}'''
        obj_flow_objects_actions_child_content = obj_flow_action_child_header + obj_flow_action_child_definition
        obj_flow_objects_actions_content += obj_flow_objects_actions_child_content
        var_x_pos += 100
        var_y_pos += 50
    current_time = str(current_time)
    obj_flow_objects_actions_content = obj_flow_objects_actions_content.replace("}{","},{")
    obj_array_assets = []
    obj_list_assets = []
    if ctid_assets is None:
        obj_list_assets = input("Give a single or a semicolon separated list of asset names to generate (for example: SYSTEM01;SRV-EXCH-01;Obsolete Device). Simply press enter if no assets need to be defined: ")
        if not obj_list_assets:
            pass
        else:
            obj_list_assets = obj_list_assets.split(';')
    else:
        obj_list_assets = ctid_assets.split(";")
    obj_array_assets = [{'asset': asset_name} for asset_name in obj_list_assets]
    asset_flow_object_child_GUID_list = []
    for asset in obj_array_assets:
        asset_guid = str(uuid.uuid4())
        asset_flow_object_child_GUID_list.append(asset_guid)
    obj_flow_objects_asset_content = ""
    asset_flow_action_children_definition = ""
    asset_flow_action_children_header = ""
    asset_flow_action_complete = ""
    var_x_pos = 100
    var_y_pos = -300
    range_techniques = range(len(obj_array_assets))
    for technique in range_techniques:
        obj_list_flow_objects_asset_child_GUID = []
        for count in range(1, 13):
            obj_action_child_GUID = str(uuid.uuid4())
            obj_list_flow_objects_asset_child_GUID.append(str(obj_action_child_GUID))
        asset_flow_action_child_GUID_group = '","'.join(obj_list_flow_objects_asset_child_GUID)
        asset_flow_action_children_header = f'{{"id":"{asset_flow_object_child_GUID_list[technique]}","x":{var_x_pos},"y":{var_y_pos},"attrs":256,"template":"asset","children":["{asset_flow_action_child_GUID_group}"],"properties":[["name","{obj_array_assets[technique]["asset"]}"],["description","DESCRIPTION_PLACEHOLDER"]]}},'
        asset_flow_action_children_definition = f'''{{"id":"{obj_list_flow_objects_asset_child_GUID[0]}","x":0,"y":0,"attrs":0,"template":"@__builtin__anchor","children":[],"properties":[],"angle":1}},{{"id":"{obj_list_flow_objects_asset_child_GUID[1]}","x":0,"y":0,"attrs":0,"template":"@__builtin__anchor","children":[],"properties":[],"angle":1}},{{"id":"{obj_list_flow_objects_asset_child_GUID[2]}","x":0,"y":0,"attrs":0,"template":"@__builtin__anchor","children":[],"properties":[],"angle":1}},{{"id":"{obj_list_flow_objects_asset_child_GUID[3]}","x":0,"y":0,"attrs":0,"template":"@__builtin__anchor","children":[],"properties":[],"angle":0}},{{"id":"{obj_list_flow_objects_asset_child_GUID[4]}","x":0,"y":0,"attrs":0,"template":"@__builtin__anchor","children":[],"properties":[],"angle":0}},{{"id":"{obj_list_flow_objects_asset_child_GUID[5]}","x":0,"y":0,"attrs":0,"template":"@__builtin__anchor","children":[],"properties":[],"angle":0}},{{"id":"{obj_list_flow_objects_asset_child_GUID[6]}","x":0,"y":0,"attrs":0,"template":"@__builtin__anchor","children":[],"properties":[],"angle":1}},{{"id":"{obj_list_flow_objects_asset_child_GUID[7]}","x":0,"y":0,"attrs":0,"template":"@__builtin__anchor","children":[],"properties":[],"angle":1}},{{"id":"{obj_list_flow_objects_asset_child_GUID[8]}","x":0,"y":0,"attrs":0,"template":"@__builtin__anchor","children":[],"properties":[],"angle":1}},{{"id":"{obj_list_flow_objects_asset_child_GUID[9]}","x":0,"y":0,"attrs":0,"template":"@__builtin__anchor","children":[],"properties":[],"angle":0}},{{"id":"{obj_list_flow_objects_asset_child_GUID[10]}","x":0,"y":0,"attrs":0,"template":"@__builtin__anchor","children":[],"properties":[],"angle":0}},{{"id":"{obj_list_flow_objects_asset_child_GUID[11]}","x":0,"y":0,"attrs":0,"template":"@__builtin__anchor","children":[],"properties":[],"angle":0}}'''
        asset_flow_action_complete = asset_flow_action_children_header + asset_flow_action_children_definition
        obj_flow_objects_asset_content += asset_flow_action_complete
        var_x_pos = var_x_pos + 100
        var_y_pos = var_y_pos + 50
    obj_flow_objects_asset_content = obj_flow_objects_asset_content.replace("}{","},{")
    if len(obj_array_assets) == 0:
        obj_list_flow_property_GUID = '","'.join([attack['guid'] for attack in array_obj_sorted_construct])
    else:
        obj_list_flow_property_GUID = '","'.join([attack['guid'] for attack in array_obj_sorted_construct] + asset_flow_object_child_GUID_list)
    obj_flow_objects_header_0 = '''"objects":[{"id":"'''
    obj_flow_objects_header_1 = var_obj_flow_property_GUID
    obj_flow_objects_header_2 = '''","x":-290,"y":-170,"attrs":0,"template":"flow","children":["'''
    obj_flow_objects_header_3 = obj_list_flow_property_GUID
    obj_flow_objects_header_4 = '''"],"properties":[["name","'''
    obj_flow_objects_header_5 = flow_name_content
    obj_flow_objects_header_6 = '''"],["description",null],["author",[["name","CPIRT"],["identity_class","db0f6f37ebeb6ea09489124345af2a45"],["contact_information","emergency-response@checkpoint.com"]]],["scope","3e072748feb6ecd1b1ba397704e009c0"],["external_references",[]],["created","'''
    obj_flow_objects_header_7 = current_time
    obj_flow_objects_header_8 = '''"]]},'''
    obj_flow_objects_header = obj_flow_objects_header_0 + obj_flow_objects_header_1 + obj_flow_objects_header_2 + obj_flow_objects_header_3 + obj_flow_objects_header_4 + obj_flow_objects_header_5 + obj_flow_objects_header_6 + obj_flow_objects_header_7 + obj_flow_objects_header_8
    obj_flow_template_footer = '''],"location":{"x":-0.5,"y":-0.5,"k":1}}'''
    if len(obj_array_assets) == 0:
        obj_flow = obj_flow_template_header + obj_flow_objects_header + obj_flow_objects_actions_content + obj_flow_template_footer
    else:
        obj_flow = obj_flow_template_header + obj_flow_objects_header + obj_flow_objects_actions_content + "," + obj_flow_objects_asset_content + obj_flow_template_footer
    with open(file_afb_ctid_flow, "w") as file_flow:
        file_flow.write(obj_flow)

def get_document_prefix(prefix):
    if not prefix:
        document_prefix_content = input("Provide the prefix of the generated documents. This could be the case number or name. This will also be used to name the Navigator Layer and CTID Flow. Simply press enter if none is required.")
        if document_prefix_content:
            document_prefix = (document_prefix_content.lower()).replace(" ","_") + "_"
        else:
            document_prefix_content = "Untitled"
            document_prefix = ""
    else:
        document_prefix_content = prefix
        document_prefix = (prefix.lower()).replace(" ","_") + "_"
    globals()["document_prefix_content"] = document_prefix_content
    globals()["document_prefix"] = document_prefix

def new_attackrecommendations(prefix=None,ciscontrols=True,nistcontrols=False):
    get_document_prefix(prefix)
    new_attackdocintroduction()
    new_attackdocmitigations(ciscontrols,nistcontrols)
    new_attackdocdetections()
    new_attackdocvalidations()

def new_attacksighting():
    sightings_id = str(uuid.uuid4())
    file_sighting_json = sightings_id + "_sighting.json"
    file_sighting_json = os.path.join(case_path, file_sighting_json)
    file_json_sighting_template = os.path.join(template_path, "sightings_template.json")
    sighting_version = "2.0"
    detection_type = "human_validated"
    sightings_techniques_array = []
    sightings_techniques_array = [attack['attack_id'] for attack in array_obj_sorted_construct]
    format = "%Y-%m-%dT%H:%M:%SZ"
    while True:
        sighting_start = input("\u2139 Please provide the start time for the sightings.\n\u2328 Please use RFC 3339 timestamps in UTC time [2022-12-22T12:03:23Z]: ")
        try:
            parseddate = datetime.strptime(sighting_start, format)
            break
        except ValueError:
            print("\u26A0 Invalid input. Try again.")
    naics_list = {11: "Agriculture, Forestry, Fishing and Hunting",
                 21: "Mining, Quarrying, and Oil and Gas Extraction",
                 22: "Utilities",
                 23: "Construction",
                 31: "Manufacturing",
                 32: "Manufacturing",
                 33: "Manufacturing",
                 42: "Wholesale Trade",
                 44: "Retail Trade",
                 45: "Retail Trade",
                 48: "Transportation and Warehousing",
                 49: "Transportation and Warehousing",
                 51: "Information",
                 52: "Finance and Insurance",
                 53: "Real Estate and Rental and Leasing",
                 54: "Professional, Scientific, and Technical Services",
                 55: "Management of Companies and Enterprises",
                 56: "Administrative and Support and Waste Management and Remediation Services",
                 61: "Educational Services",
                 62: "Health Care and Social Assistance",
                 71: "Arts, Entertainment, and Recreation",
                 72: "Accommodation and Food Services",
                 81: "Other Services (except Public Administration)",
                 92: "Public Administration"}
    while True:
        victim_sector = input("\u2328 Provide the victim sector NAICS code, first 2 digits only [eg 22]. Tap Enter to present the list: ")
        try:
            if int(victim_sector) not in naics_list.keys():
                raise ValueError
            else:
                victim_sector_name = naics_list[int(victim_sector)]
                print("\u2705 You selected the following sector:", victim_sector_name)
                break
        except ValueError as error:
            print("\u26A0", victim_sector, "is not in the NAICS list. Verify your input please.")
            naics_table = sorted(naics_list.items())
            print("Refer to the following list:")
            for sector, name in naics_table:
                print(sector, "-", name)
    iso_country_list = {"AF": "The Islamic Republic of Afghanistan","AX": "Åland","AL": "The Republic of Albania","DZ": "The People's Democratic Republic of Algeria","AS": "The Territory of American Samoa","AD": "The Principality of Andorra","AO": "The Republic of Angola","AI": "Anguilla","AQ": "All land and ice shelves south of the 60th parallel south","AG": "Antigua and Barbuda","AR": "The Argentine Republic","AM": "The Republic of Armenia","AW": "Aruba","AU": "The Commonwealth of Australia","AT": "The Republic of Austria","AZ": "The Republic of Azerbaijan","BS": "The Commonwealth of The Bahamas","BH": "The Kingdom of Bahrain","BD": "The People's Republic of Bangladesh","BB": "Barbados","BY": "The Republic of Belarus","BE": "The Kingdom of Belgium","BZ": "Belize","BJ": "The Republic of Benin","BM": "Bermuda","BT": "The Kingdom of Bhutan","BO": "The Plurinational State of Bolivia","BQ": "Bonaire, Sint Eustatius and Saba","BA": "Bosnia and Herzegovina","BW": "The Republic of Botswana","BV": "Bouvet Island","BR": "The Federative Republic of Brazil","IO": "The British Indian Ocean Territory","BN": "The Nation of Brunei, the Abode of Peace","BG": "The Republic of Bulgaria","BF": "Burkina Faso","BI": "The Republic of Burundi","CV": "The Republic of Cabo Verde","KH": "The Kingdom of Cambodia","CM": "The Republic of Cameroon","CA": "Canada","KY": "The Cayman Islands","CF": "The Central African Republic","TD": "The Republic of Chad","CL": "The Republic of Chile","CN": "The People's Republic of China","CX": "The Territory of Christmas Island","CC": "The Territory of Cocos (Keeling) Islands","CO": "The Republic of Colombia","KM": "The Union of the Comoros","CD": "The Democratic Republic of the Congo","CG": "The Republic of the Congo","CK": "The Cook Islands","CR": "The Republic of Costa Rica","CI": "The Republic of Côte d'Ivoire","HR": "The Republic of Croatia","CU": "The Republic of Cuba","CW": "The Country of Curaçao","CY": "The Republic of Cyprus","CZ": "The Czech Republic","DK": "The Kingdom of Denmark","DJ": "The Republic of Djibouti","DM": "The Commonwealth of Dominica","DO": "The Dominican Republic","EC": "The Republic of Ecuador","EG": "The Arab Republic of Egypt","SV": "The Republic of El Salvador","GQ": "The Republic of Equatorial Guinea","ER": "The State of Eritrea","EE": "The Republic of Estonia","SZ": "The Kingdom of Eswatini","ET": "The Federal Democratic Republic of Ethiopia","FK": "The Falkland Islands","FO": "The Faroe Islands","FJ": "The Republic of Fiji","FI": "The Republic of Finland","FR": "The French Republic","GF": "Guyane","PF": "French Polynesia","TF": "The French Southern and Antarctic Lands","GA": "The Gabonese Republic","GM": "The Republic of The Gambia","GE": "Georgia","DE": "The Federal Republic of Germany","GH": "The Republic of Ghana","GI": "Gibraltar","GR": "The Hellenic Republic","GL": "Kalaallit Nunaat","GD": "Grenada","GP": "Guadeloupe","GU": "The Territory of Guam","GT": "The Republic of Guatemala","GG": "The Bailiwick of Guernsey","GN": "The Republic of Guinea","GW": "The Republic of Guinea-Bissau","GY": "The Co-operative Republic of Guyana","HT": "The Republic of Haiti","HM": "The Territory of Heard Island and McDonald Islands","VA": "The Holy See","HN": "The Republic of Honduras","HK": "The Hong Kong Special Administrative Region of China[10]","HU": "Hungary","IS": "Iceland","IN": "The Republic of India","ID": "The Republic of Indonesia","IR": "The Islamic Republic of Iran","IQ": "The Republic of Iraq","IE": "Ireland","IM": "The Isle of Man","IL": "The State of Israel","IT": "The Italian Republic","JM": "Jamaica","JP": "Japan","JE": "The Bailiwick of Jersey","JO": "The Hashemite Kingdom of Jordan","KZ": "The Republic of Kazakhstan","KE": "The Republic of Kenya","KI": "The Republic of Kiribati","KP": "The Democratic People's Republic of Korea","KR": "The Republic of Korea","KW": "The State of Kuwait","KG": "The Kyrgyz Republic","LA": "The Lao People's Democratic Republic","LV": "The Republic of Latvia","LB": "The Lebanese Republic","LS": "The Kingdom of Lesotho","LR": "The Republic of Liberia","LY": "The State of Libya","LI": "The Principality of Liechtenstein","LT": "The Republic of Lithuania","LU": "The Grand Duchy of Luxembourg","MO": "The Macao Special Administrative Region of China[11]","MK": "The Republic of North Macedonia[12]","MG": "The Republic of Madagascar","MW": "The Republic of Malawi","MY": "Malaysia","MV": "The Republic of Maldives","ML": "The Republic of Mali","MT": "The Republic of Malta","MH": "The Republic of the Marshall Islands","MQ": "Martinique","MR": "The Islamic Republic of Mauritania","MU": "The Republic of Mauritius","YT": "The Department of Mayotte","MX": "The United Mexican States","FM": "The Federated States of Micronesia","MD": "The Republic of Moldova","MC": "The Principality of Monaco","MN": "Mongolia","ME": "Montenegro","MS": "Montserrat","MA": "The Kingdom of Morocco","MZ": "The Republic of Mozambique","MM": "The Republic of the Union of Myanmar","NA": "The Republic of Namibia","NR": "The Republic of Nauru","NP": "The Federal Democratic Republic of Nepal","NL": "The Kingdom of the Netherlands","NC": "New Caledonia","NZ": "New Zealand","NI": "The Republic of Nicaragua","NE": "The Republic of the Niger","NG": "The Federal Republic of Nigeria","NU": "Niue","NF": "The Territory of Norfolk Island","MP": "The Commonwealth of the Northern Mariana Islands","NO": "The Kingdom of Norway","OM": "The Sultanate of Oman","PK": "The Islamic Republic of Pakistan","PW": "The Republic of Palau","PS": "The State of Palestine","PA": "The Republic of Panamá","PG": "The Independent State of Papua New Guinea","PY": "The Republic of Paraguay","PE": "The Republic of Perú","PH": "The Republic of the Philippines","PN": "The Pitcairn, Henderson, Ducie and Oeno Islands","PL": "The Republic of Poland","PT": "The Portuguese Republic","PR": "The Commonwealth of Puerto Rico","QA": "The State of Qatar","RE": "Réunion","RO": "Romania","RU": "The Russian Federation","RW": "The Republic of Rwanda","BL": "The Collectivity of Saint-Barthélemy","SH": "Saint Helena, Ascension and Tristan da Cunha","KN": "Saint Kitts and Nevis","LC": "Saint Lucia","MF": "The Collectivity of Saint-Martin","PM": "The Overseas Collectivity of Saint-Pierre and Miquelon","VC": "Saint Vincent and the Grenadines","WS": "The Independent State of Samoa","SM": "The Republic of San Marino","ST": "The Democratic Republic of São Tomé and Príncipe","SA": "The Kingdom of Saudi Arabia","SN": "The Republic of Senegal","RS": "The Republic of Serbia","SC": "The Republic of Seychelles","SL": "The Republic of Sierra Leone","SG": "The Republic of Singapore","SX": "Sint Maarten","SK": "The Slovak Republic","SI": "The Republic of Slovenia","SB": "The Solomon Islands","SO": "The Federal Republic of Somalia","ZA": "The Republic of South Africa","GS": "South Georgia and the South Sandwich Islands","SS": "The Republic of South Sudan","ES": "The Kingdom of Spain","LK": "The Democratic Socialist Republic of Sri Lanka","SD": "The Republic of the Sudan","SR": "The Republic of Suriname","SJ": "Svalbard and Jan Mayen","SE": "The Kingdom of Sweden","CH": "The Swiss Confederation","SY": "The Syrian Arab Republic","TW": "The Republic of China","TJ": "The Republic of Tajikistan","TZ": "The United Republic of Tanzania","TH": "The Kingdom of Thailand","TL": "The Democratic Republic of Timor-Leste","TG": "The Togolese Republic","TK": "Tokelau","TO": "The Kingdom of Tonga","TT": "The Republic of Trinidad and Tobago","TN": "The Republic of Tunisia","TR": "The Republic of Türkiye","TM": "Turkmenistan","TC": "The Turks and Caicos Islands","TV": "Tuvalu","UG": "The Republic of Uganda","UA": "Ukraine","AE": "The United Arab Emirates","GB": "The United Kingdom of Great Britain and Northern Ireland","UM": "Baker Island, Howland Island, Jarvis Island, Johnston Atoll, Kingman Reef, Midway Atoll, Navassa Island, Palmyra Atoll, and Wake Island","US": "The United States of America","UY": "The Oriental Republic of Uruguay","UZ": "The Republic of Uzbekistan","VU": "The Republic of Vanuatu","VE": "The Bolivarian Republic of Venezuela","VN": "The Socialist Republic of Viet Nam","VG": "The Virgin Islands","VI": "The Virgin Islands of the United States","WF": "The Territory of the Wallis and Futuna Islands","EH": "The Sahrawi Arab Democratic Republic","YE": "The Republic of Yemen","ZM": "The Republic of Zambia","ZW": "The Republic of Zimbabwe"}
    while True:
        victim_country = input("\u2328 Provide the victim ISO 3166-1 alpha-2 country code [eg BE]: ")
        victim_country = victim_country.upper()
        try:
            if (victim_country) not in iso_country_list.keys():
                raise ValueError
            else:
                victim_country_name = iso_country_list[(victim_country)]
                print("\u2705 You selected the following country:", victim_country_name)
                break
        except ValueError as error:
            print("\u26A0", victim_country, "is not in the ISO Country list. Verify your input please.")
    detection_list = ["host_based", "network_based", "cloud_based"]
    while True:
        detection_source = input("\u2328 Define the detection source [host_based, network_based, cloud_based]: ")
        detection_source = detection_source.lower()
        try:
            index = detection_list.index(detection_source)
            break
        except:
            print("\u26A0", detection_source, "is not in the list. Verify your input please.")
    platform_list = ["windows","macos","nix","other"]
    while True:
        victim_platform_env = input("\u2328 Define the platform [windows, macos, nix, other]: ")
        victim_platform_env = victim_platform_env.lower()
        try:
            index = platform_list.index(victim_platform_env)
            break
        except ValueError:
            print("\u26A0", victim_platform_env, "is not in the list. Verify your input please.")
    privilege_list = ["system","admin","user","none"]
    while True:
        victim_privilegelevel = input("\u2328 Provide the privilege level [system, admin, user, none]: ")
        victim_privilegelevel = victim_privilegelevel.lower()
        try:
            index = privilege_list.index(victim_privilegelevel)
            break
        except ValueError:
            print("\u26A0", victim_privilegelevel, "is not in the list. Verify your input please.")
    sighting_software = input("Provide the malicious software name that was observed. This should be an exact name from the list https://attack.mitre.org/software/. Simply press enter if not applicable.")
    with open(file_json_sighting_template,'r+') as file:
        sightings_array_json = json.load(file)
    sightings_array_json["version"] = sighting_version
    sightings_array_json["id"] = sightings_id
    sightings_array_json["start_time"] = sighting_start
    sightings_array_json["tid"] += sightings_techniques_array
    sightings_array_json["detection_type"] = detection_type
    sightings_array_json["detection_source"] = detection_source
    sightings_array_json["sector"] = victim_sector
    sightings_array_json["country"] = victim_country
    sightings_array_json["platform"] = victim_platform_env
    sightings_array_json["privilege_level"] = victim_privilegelevel
    if sighting_software:
        sightings_array_json["software_name"] = sighting_software
    with open(file_sighting_json, 'w') as file_sighting:
        file_sighting.write(json.dumps(sightings_array_json, indent=4))

def new_attacknavigatorlayer():
    file_prefix = ""
    file_navigator_layer_json = os.path.join(case_path, file_prefix, document_prefix + "navigator_layer.json")
    file_json_navigator_layer_template = os.path.join(template_path, "navigator_template.json")
    var_obj_layer_technique_property_colour = "#c41a9f"
    var_obj_layer_tactic_property_colour = "#c41a9f"
    array_obj_navigator_techniques = []
    for attack in array_obj_sorted_construct:
            techniqueID = attack['attack_id']
            tactic = attack['attack_tactics']
            color = var_obj_layer_technique_property_colour
            new_obj = {
                'techniqueID': techniqueID,
                'tactic': tactic,
                'color': color
            }
            array_obj_navigator_techniques.append(new_obj)
    with open(file_json_navigator_layer_template,'r+') as file:
            obj_complete_navigator_layer = json.load(file)
    obj_complete_navigator_layer["name"] = document_prefix_content
    obj_complete_navigator_layer["tacticRowBackground"] = var_obj_layer_tactic_property_colour
    obj_complete_navigator_layer["techniques"] += array_obj_navigator_techniques
    with open(file_navigator_layer_json, "w") as file_navigator_layer:
        file_navigator_layer.write(json.dumps(obj_complete_navigator_layer, indent=4))
